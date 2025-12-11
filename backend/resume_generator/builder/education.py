from typing import Dict
import re

from .common import replace_placeholders_in_doc
from docx import Document
from docx.oxml.table import CT_Tc


def italicize_runs(runs, text_to_match: str) -> None:
    """Italicize any run containing the given text."""
    if not text_to_match:
        return
    for r in runs:
        if r.text and text_to_match in r.text:
            r.font.italic = True

def render_education(doc, ctx) -> None:
    """
    Render primary education placeholders for single-entry templates.
    """
    field_values: Dict[str, str] = {
        "edu_name": ctx.edu_name or "",
        "edu_degree": ctx.edu_degree or "",
        "edu_minor": ctx.edu_minor or "",
        "edu_location": f"( {ctx.edu_location} )" if ctx.edu_location else "",
        "edu_gpa": ctx.edu_gpa or "",
        "edu_date": ctx.edu_date or "",
        "edu_honors": ctx.edu_honors or "",
        "edu_clubs": ctx.edu_clubs or "",
        "edu_coursework": ctx.edu_coursework or "",
    }

    # fill in placeholders in document.
    replace_placeholders_in_doc(doc, field_values)

    # handle inline conditionals `{? field:INLINE}` for education across all paragraphs (including tables)
    _process_inline(doc, "edu_location", field_values["edu_location"])
    _process_inline(doc, "edu_date", field_values["edu_date"])
    _process_inline(doc, "edu_gpa", field_values["edu_gpa"])
    _process_inline(doc, "edu_honors", field_values["edu_honors"])
    _process_inline(doc, "edu_clubs", field_values["edu_clubs"])
    _process_inline(doc, "edu_coursework", field_values["edu_coursework"])
    _process_inline(doc, "edu_minor", field_values["edu_minor"])


INLINE_PATTERN_CACHE: Dict[str, re.Pattern] = {}


def _process_inline(doc: Document, field: str, value: str) -> None:
    """
    Handle `{? field:INLINE}` markers across paragraphs and table paragraphs:
      - if value is truthy: remove the marker text, preserving run formatting
      - if value is falsy: remove only the marker and any immediately preceding separator (comma, dash, etc.)
    """
    if field not in INLINE_PATTERN_CACHE:
        INLINE_PATTERN_CACHE[field] = re.compile(
            r"\{\?\s*" + re.escape(field) + r"\s*:\s*INLINE\s*\}", re.IGNORECASE
        )
    pat = INLINE_PATTERN_CACHE[field]
    
    # Pattern to match separators before the marker (comma, dash, space combinations)
    separator_pat = re.compile(r'[\s,–—\-]+\{\?\s*' + re.escape(field) + r'\s*:\s*INLINE\s*\}', re.IGNORECASE)

    def clean_text(text: str) -> str:
        t = separator_pat.sub("", text)
        t = pat.sub("", t)
        fl = field.lower()
        if fl == "edu_gpa":
            t = re.sub(r"\bGPA\b[:\-\s]*", "", t, flags=re.IGNORECASE)
        elif fl == "edu_minor":
            t = re.sub(r"\s*[–—-]\s*Minor\s+in\s*$", "", t, flags=re.IGNORECASE)
            t = re.sub(r"\s*[–—-]\s*Minor\s+in\s+", " ", t, flags=re.IGNORECASE)
        elif fl == "edu_honors":
            t = re.sub(r"\bHonors\s*&\s*Awards:?\s*", "", t, flags=re.IGNORECASE)
        elif fl == "edu_clubs":
            t = re.sub(r"\bClubs\s*&\s*Extracurriculars:?\s*", "", t, flags=re.IGNORECASE)
        elif fl == "edu_coursework":
            t = re.sub(r"\bRelevant\s*Coursework:?\s*", "", t, flags=re.IGNORECASE)
        t = re.sub(r"\s{2,}", " ", t).strip()
        return t

    def _process_para(para):
        combined = "".join(run.text for run in para.runs)

        if not pat.search(combined):
            return

        if value:
            for run in para.runs:
                if run.text:
                    run.text = pat.sub("", run.text)
        else:
            for run in para.runs:
                if run.text:
                    run.text = clean_text(run.text)

        combined_after = "".join(run.text for run in para.runs).strip()

        if combined_after == "":
            _safe_remove_paragraph(para)

    for para in doc.paragraphs:
        _process_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_para(para)


# this is how we safely remove a paragraph from a document if its empty after toggle.
def _safe_remove_paragraph(para):
    p = para._element
    parent = p.getparent()

    if parent is None:
        return

    # CASE #1 : inside a table cell - do NOT remove, just collapse
    if isinstance(parent, CT_Tc):
        cell_paras = parent.findall(".//w:p", parent.nsmap)

        if len(cell_paras) > 1:
            try:
                parent.remove(p)
            except Exception:
                para.text = ""
        else:
            # LAST paragraph -> collapse instead of removing
            para.text = ""
            
            for run in para.runs:
                try:
                    run._element.clear_content()
                except Exception:
                    run.text = ""
            
            fmt = para.paragraph_format
            fmt.space_before = 0
            fmt.space_after = 0
            fmt.left_indent = 0
            fmt.right_indent = 0
            fmt.first_line_indent = 0
            fmt.line_spacing = 1
        
        return
    
    # CASE #2: normal paragraph
    try:
        parent.remove(p)
    except:
        para.text = ""

