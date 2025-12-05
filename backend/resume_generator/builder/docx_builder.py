"""
simple, template‑driven docx resume builder (v1 – header only).

placeholder language (current scope)
------------------------------------

supported now:
  {field}
    - simple inline replacement
    - case‑insensitive: {name}, {NAME}, {NaMe} all map to "name"
    - used in header:
        {name}
        {email}
        {github}
        {linkedin}
        {portfolio}

conditional markers ({?field}, {?field:INLINE}) and section builders
will be added back later; for now we only render the header and focus on
styling preservation.

key ideas:
  - templates own all visual styling
  - builder only replaces placeholders in paragraphs / table cells
  - run‑level formatting is preserved whenever possible
"""

import os
import re
from typing import Dict, List, Tuple, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..context import ResumeRenderContext, build_resume_render_context


# public entry point

def build_docx(
    user,
    template_path: str,
    overrides: Optional[Dict[str, str]] = None,
    margin_overrides: Optional[Dict[str, float]] = None,
    header_line: Optional[str] = None,
    header_alignment: Optional[str] = None,
) -> Document:

    doc = Document(template_path)

    # apply margin overrides to the template if provided
    if margin_overrides:
        for section in doc.sections:
            if "margin_top" in margin_overrides and margin_overrides["margin_top"] is not None:
                section.top_margin = Inches(margin_overrides["margin_top"])
            if "margin_bottom" in margin_overrides and margin_overrides["margin_bottom"] is not None:
                section.bottom_margin = Inches(margin_overrides["margin_bottom"])
            if "margin_left" in margin_overrides and margin_overrides["margin_left"] is not None:
                section.left_margin = Inches(margin_overrides["margin_left"])
            if "margin_right" in margin_overrides and margin_overrides["margin_right"] is not None:
                section.right_margin = Inches(margin_overrides["margin_right"])

    # 1) build render context
    ctx: ResumeRenderContext = build_resume_render_context(user)

    # apply simple overrides (header fields) if provided
    if overrides:
        for key, value in overrides.items():
            if hasattr(ctx, key):
                setattr(ctx, key, value)

    # apply assembled header line if provided
    if header_line is not None:
        ctx.header_line = header_line

    # 2) render header only (global simple placeholders)
    _render_header(doc, ctx)

    # 3) apply header alignment after placeholders are replaced
    if header_alignment:
        _apply_header_alignment(doc, ctx.header_line or "", header_alignment)

    # Later we will call:
    #   _render_education(...)
    #   _render_experience(...)
    #   _render_projects(...)
    #   _render_skills(...)
    # as we grow this file again.

    return doc


# header rendering

def _render_header(doc: Document, ctx: ResumeRenderContext) -> None:
    field_values: Dict[str, str] = {
        "name": ctx.name or "",
        "email": ctx.email or "",
        "github": ctx.github or "",
        "linkedin": ctx.linkedin or "",
        "portfolio": ctx.portfolio or "",
        "phone": ctx.phone or "",
        "location": ctx.location or "",
        "header_line": ctx.header_line or "",
    }

    _replace_placeholders_in_doc(doc, field_values)
    _ensure_header_line_spacing(doc)


# core placeholder engine (inline, no conditionals yet)


def _collapse_pipe_separators(text: str) -> str:
    """
    Collapse pipe-delimited segments, removing empties and trimming edges.
    Example: " |  a  |   | b |  | " -> "a | b"
    """
    parts = [p.strip() for p in text.split("|") if p.strip()]
    return " | ".join(parts)


def _replace_placeholders_in_doc(doc: Document, field_values: Dict[str, str]) -> None:
    if not field_values:
        return

    # compile patterns once per field
    patterns: Dict[str, re.Pattern] = {
        field: re.compile(r"\{\s*" + re.escape(field) + r"\s*\}", re.IGNORECASE)
        for field in field_values.keys()
    }

    def _apply_all(text: str) -> str:
        """apply all {field} replacements to a single string."""
        new_text = text
        for field, pattern in patterns.items():
            value = field_values.get(field) or ""
            if pattern.search(new_text):
                new_text = pattern.sub(value, new_text)
        return new_text

    def _replace_in_paragraph(paragraph) -> None:
        # if there is nothing here, skip quickly
        if not paragraph.runs and not paragraph.text:
            return

        full_text = (
            "".join(run.text for run in paragraph.runs)
            if paragraph.runs
            else paragraph.text
        )

        if not full_text:
            return

        # fast path – check if there is any placeholder at all
        if "{" not in full_text or "}" not in full_text:
            return

        # pass 1: run-level replacement (no restructuring)
        if paragraph.runs:
            for run in paragraph.runs:
                original = run.text
                if not original:
                    continue
                replaced = _apply_all(original)
                if replaced != original:
                    run.text = replaced

            # after run-level replacements, recompute full_text to see if any placeholders remain
            full_text = "".join(run.text for run in paragraph.runs)

        # pass 2: multi-run placeholders (structured rebuild) only if placeholders remain
        new_full_text = _apply_all(full_text)
        if new_full_text == full_text:
            # no placeholders left in this paragraph
            # cleanup separators if needed
            final_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text or ""
            collapsed = _collapse_pipe_separators(final_text)
            if collapsed != final_text:
                paragraph.text = collapsed
            return

        # if there are no runs, we cannot preserve run formatting anyway
        # so we just overwrite the paragraph text
        if not paragraph.runs:
            paragraph.text = _collapse_pipe_separators(new_full_text)
            return

        runs = paragraph.runs
        run_texts = [r.text for r in runs]

        # build a map from character offset in full_text to (run_index, char_index)
        char_map: List[Tuple[int, int]] = []
        for i, txt in enumerate(run_texts):
            for j in range(len(txt)):
                char_map.append((i, j))

        if len(char_map) != len(full_text):
            # something is off – fall back to a simpler replacement
            paragraph.text = new_full_text
            return

        # find all placeholder spans in the original full_text
        spans: List[Tuple[int, int, str]] = []  # (start_idx, end_idx, replacement_text)
        for field, pattern in patterns.items():
            value = field_values.get(field) or ""
            for m in pattern.finditer(full_text):
                spans.append((m.start(), m.end(), value))

        if not spans:
            return

        spans.sort(key=lambda x: x[0])

        # rebuild segments (text + font)
        segments: List[Tuple[str, object]] = []
        cursor = 0

        for start, end, replacement in spans:
            # text before placeholder
            if cursor < start:
                segment_text = full_text[cursor:start]
                run_idx, _ = char_map[cursor]
                segments.append((segment_text, runs[run_idx].font))

            # replacement text styled like the run where placeholder starts
            run_idx, _ = char_map[start]
            segments.append((replacement, runs[run_idx].font))

            cursor = end

        # trailing text after last placeholder
        if cursor < len(full_text):
            segment_text = full_text[cursor:]
            run_idx, _ = char_map[cursor]
            segments.append((segment_text, runs[run_idx].font))

        # clear and rebuild runs
        paragraph.clear()
        for text, font_obj in segments:
            if not text:
                continue
            new_run = paragraph.add_run(text)
            try:
                new_run.font.size = font_obj.size
                new_run.font.bold = font_obj.bold
                new_run.font.italic = font_obj.italic
                if font_obj.color and font_obj.color.rgb:
                    new_run.font.color.rgb = font_obj.color.rgb
            except Exception:
                # best effort on formatting
                pass

        # final cleanup of separators for rebuilt paragraphs
        final_text = paragraph.text
        collapsed = _collapse_pipe_separators(final_text)
        if collapsed != final_text:
            paragraph.text = collapsed
        return

    # apply to normal paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para)

    # apply inside table cells (headers are often built with tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)


def _ensure_header_line_spacing(doc: Document) -> None:
    """
    Add a small bottom padding to paragraphs that contain header_line content
    to avoid the next line crowding when it wraps. Uses paragraph spacing
    instead of manual newlines so it won't be overwritten by placeholder replacement.
    """
    target_keys = {"{header_line}", "{ header_line }"}
    for para in doc.paragraphs:
        text = para.text or ""
        if any(k in text for k in target_keys):
            if para.paragraph_format and para.paragraph_format.space_after is None:
                para.paragraph_format.space_after = Pt(6)  # ~0.08"


def _apply_header_alignment(doc: Document, header_text: str, alignment: str) -> None:
    """
    Apply alignment to paragraphs that carry name or header_line after replacement.
    alignment: 'left' | 'center' | 'right'
    """
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    target_align = align_map.get(alignment.lower())
    if target_align is None:
        return

    target_keys = {"{name}", "{header_line}", "{ header_line }"}
    header_text = (header_text or "").strip()

    def _maybe_align(paragraph) -> None:
        text = (paragraph.text or "").strip()
        if not text:
            return
        # match if the original placeholder markers are still present
        if any(k in text for k in target_keys):
            if paragraph.paragraph_format:
                paragraph.alignment = target_align
            return
        # match if the rendered text equals header_text or contains it
        if header_text and (text == header_text or header_text in text):
            if paragraph.paragraph_format:
                paragraph.alignment = target_align
            return
        # match if this is the name line
        if text and text == (header_text or "").split(" | ")[0]:
            if paragraph.paragraph_format:
                paragraph.alignment = target_align
            return

    for para in doc.paragraphs:
        _maybe_align(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _maybe_align(para)
