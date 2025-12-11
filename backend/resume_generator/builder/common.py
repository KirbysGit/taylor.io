import re
from typing import Dict, List, Tuple, Optional
from copy import deepcopy

from docx import Document
from docx.shared import Inches
from docx.text.run import Run
from docx.enum.text import WD_ALIGN_PARAGRAPH


def apply_margins(doc: Document, margin_overrides: Optional[Dict[str, float]]) -> None:
    if not margin_overrides:
        return
    for section in doc.sections:
        if margin_overrides.get("margin_top") is not None:
            section.top_margin = Inches(margin_overrides["margin_top"])
        if margin_overrides.get("margin_bottom") is not None:
            section.bottom_margin = Inches(margin_overrides["margin_bottom"])
        if margin_overrides.get("margin_left") is not None:
            section.left_margin = Inches(margin_overrides["margin_left"])
        if margin_overrides.get("margin_right") is not None:
            section.right_margin = Inches(margin_overrides["margin_right"])


def collapse_pipe_separators(text: str) -> str:
    parts = [p.strip() for p in text.split("|") if p.strip()]
    return " | ".join(parts)


def replace_placeholders_in_doc(doc: Document, field_values: Dict[str, str]) -> None:
    if not field_values:
        return

    patterns: Dict[str, re.Pattern] = {
        field: re.compile(r"\{\s*" + re.escape(field) + r"\s*\}", re.IGNORECASE)
        for field in field_values.keys()
    }

    def _apply_all(text: str) -> str:
        new_text = text
        for field, pattern in patterns.items():
            value = field_values.get(field) or ""
            if pattern.search(new_text):
                new_text = pattern.sub(value, new_text)
        return new_text

    def _replace_in_paragraph(paragraph) -> None:
        if not paragraph.runs and not paragraph.text:
            return

        full_text = (
            "".join(run.text for run in paragraph.runs)
            if paragraph.runs
            else paragraph.text
        )
        if not full_text:
            return

        if "{" not in full_text or "}" not in full_text:
            return

        if paragraph.runs:
            for run in paragraph.runs:
                original = run.text
                if not original:
                    continue
                replaced = _apply_all(original)
                if replaced != original:
                    run.text = replaced
            full_text = "".join(run.text for run in paragraph.runs)

        new_full_text = _apply_all(full_text)
        if new_full_text == full_text:
            final_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text or ""
            collapsed = collapse_pipe_separators(final_text)
            if collapsed != final_text:
                paragraph.text = collapsed
            return

        if not paragraph.runs:
            paragraph.text = collapse_pipe_separators(new_full_text)
            return

        runs = paragraph.runs
        run_texts = [r.text for r in runs]

        char_map: List[Tuple[int, int]] = []
        for i, txt in enumerate(run_texts):
            for j in range(len(txt)):
                char_map.append((i, j))

        if len(char_map) != len(full_text):
            paragraph.text = new_full_text
            return

        spans: List[Tuple[int, int, str]] = []
        for field, pattern in patterns.items():
            value = field_values.get(field) or ""
            for m in pattern.finditer(full_text):
                spans.append((m.start(), m.end(), value))
        if not spans:
            return
        spans.sort(key=lambda x: x[0])

        segments: List[Tuple[str, Run]] = []
        cursor = 0
        for start, end, replacement in spans:
            if cursor < start:
                segment_text = full_text[cursor:start]
                run_idx, _ = char_map[cursor]
                segments.append((segment_text, runs[run_idx]))
            run_idx, _ = char_map[start]
            # dynamic text should inherit formatting from the placeholder run itself
            segments.append((replacement, runs[run_idx]))
            cursor = end
        if cursor < len(full_text):
            segment_text = full_text[cursor:]
            run_idx, _ = char_map[cursor]
            segments.append((segment_text, runs[run_idx]))

        paragraph.clear()
        for text, src_run in segments:
            if not text:
                continue
            new_run = paragraph.add_run(text)

            # clone full styling from the source run
            src_r = src_run._element
            new_r = new_run._element

            # remove any default rPr first
            if new_r.rPr is not None:
                new_r.remove(new_r.rPr)

            if src_r.rPr is not None:
                new_r.append(deepcopy(src_r.rPr))

    for para in doc.paragraphs:
        _replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)


def detect_header_paragraphs(doc: Document) -> List:
    header_keys = ("{name}", "{header_line}")
    result: List = []

    def _para_contains_keys(paragraph) -> bool:
        raw = "".join(t.text for t in paragraph._element.xpath(".//w:t"))
        return any(k in raw for k in header_keys)

    for para in doc.paragraphs:
        if _para_contains_keys(para):
            result.append(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _para_contains_keys(para):
                        result.append(para)
    return result


def apply_alignment_to_paragraphs(paragraphs: List, alignment: str) -> None:
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    target = align_map.get(alignment.lower())
    if target is None:
        return
    for para in paragraphs:
        para.alignment = target


def apply_font_family(doc: Document, font_family: str) -> None:
    if not font_family:
        return

    def _set_font(paragraph):
        for run in paragraph.runs:
            try:
                run.font.name = font_family
            except Exception:
                pass

    for para in doc.paragraphs:
        _set_font(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _set_font(para)

