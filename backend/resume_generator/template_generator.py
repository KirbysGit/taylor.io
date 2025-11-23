# template_generator.py

# Utility to programmatically create Word document templates.

# imports.
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_modern_template(output_path: str):
    """
    Create a modern-style resume template with placeholder tags.
    
    Args:
        output_path: Path where to save the template (e.g., "templates/modern.docx")
    """
    doc = Document()
    
    # set up page margins.
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # add header section with placeholders.
    heading = doc.add_heading("{NAME}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_format = heading.runs[0].font
    heading_format.size = Pt(24)
    heading_format.bold = True
    
    # add contact info line with placeholders.
    contact_para = doc.add_paragraph("{EMAIL} | {PHONE} | {GITHUB} | {LINKEDIN} | {PORTFOLIO}")
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para_format = contact_para.runs[0].font
    contact_para_format.size = Pt(10)
    
    # add spacing.
    doc.add_paragraph()
    
    # add section placeholders (these will be replaced/expanded by builder).
    # the builder will add actual content, but we can add styled section headers here.
    
    # add a horizontal line separator.
    _add_horizontal_line(doc)
    
    return doc


def create_classic_template(output_path: str):
    """
    Create a classic-style resume template with placeholder tags.
    
    Args:
        output_path: Path where to save the template (e.g., "templates/classic.docx")
    """
    doc = Document()
    
    # set up page margins.
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # add header with name.
    heading = doc.add_heading("{NAME}", level=1)
    heading_format = heading.runs[0].font
    heading_format.size = Pt(22)
    heading_format.bold = True
    
    # add contact info in separate lines.
    if "{EMAIL}":
        email_para = doc.add_paragraph("Email: {EMAIL}")
        email_para.runs[0].font.size = Pt(11)
    
    if "{PHONE}":
        phone_para = doc.add_paragraph("Phone: {PHONE}")
        phone_para.runs[0].font.size = Pt(11)
    
    # add spacing.
    doc.add_paragraph()
    
    return doc


def _add_horizontal_line(doc: Document):
    """Add a horizontal line separator."""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)
    
    # create a border element.
    pBdr = OxmlElement('w:pBdr')
    p.paragraph_format.element.get_or_add_pPr().append(pBdr)
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)


def generate_all_templates(template_dir: str = None):
    """
    Generate all default templates and save them to the templates directory.
    
    Args:
        template_dir: Directory to save templates (default: templates/ relative to this file)
    """
    if template_dir is None:
        template_dir = os.path.join(os.path.dirname(__file__), "templates")
    
    # ensure directory exists.
    os.makedirs(template_dir, exist_ok=True)
    
    # generate modern template.
    modern_path = os.path.join(template_dir, "modern.docx")
    modern_doc = create_modern_template(modern_path)
    modern_doc.save(modern_path)
    print(f"Created template: {modern_path}")
    
    # generate classic template.
    classic_path = os.path.join(template_dir, "classic.docx")
    classic_doc = create_classic_template(classic_path)
    classic_doc.save(classic_path)
    print(f"Created template: {classic_path}")


if __name__ == "__main__":
    # run this script to generate templates.
    generate_all_templates()

