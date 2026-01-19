
# imports.
from io import BytesIO
from pathlib import Path
from docx import Document
from typing import Dict, Any

# get abs path to templates directory.
TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

def build_header(header: Dict[str, Any]) -> str:
    fields = [
        header.get("email", ""),
        header.get("phone", ""),
        header.get("github", ""),
        header.get("linkedin", ""),
        header.get("location", ""),
        header.get("portfolio", ""),
    ]
    return " | ".join([field for field in fields if field.strip()])

def fill_template(document: Document, resume_data: Dict[str, Any]) -> Document:

    # get header data to fill in.
    header = resume_data.get("header", {})

    name = f"{header.get('first_name', '')} {header.get('last_name', '')}".strip()

    # fill header line.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = paragraph.text.replace("{name}", name)
                    paragraph.text = paragraph.text.replace("{header_line}", build_header(header))

    return document

def generate_resume(template_name: str, resume_data: Dict[str, Any]) -> bytes:

    # grab template from our template directory.
    template_path = TEMPLATES_DIR / f"{template_name}.docx"

    # load template.
    document = Document(str(template_path))

    # fill template based on placeholders.
    filled_document = fill_template(document, resume_data)

    # saved filled document to memory.
    return filled_document