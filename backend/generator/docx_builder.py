# docx_builder.py
# Word document generator for resumes. Uses docx_styles for template-consistent formatting.

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING, WD_UNDERLINE
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, Optional, List, Tuple

from .builders import (
    format_contact_field_display,
    format_date_month_year,
    parse_tagline_runs,
    resolve_contact_url_display,
)
from .docx_styles import get_styles, DocxStyleConfig


def _set_run_character_spacing(run, spacing_pt: float) -> None:
    """Word w:spacing on the run; val in 1/20 pt (positive = expand, negative = condense). Matches CSS letter-spacing."""
    if spacing_pt is None or spacing_pt == 0:
        return
    r_pr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(int(round(float(spacing_pt) * 20))))
    r_pr.append(spacing)


def _set_line_spacing_multiple(p, mult: float) -> None:
    """Align Word line leading with CSS line-height. Avoids Normal/defaults (often looser)."""
    m = float(mult)
    if m <= 0:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        return
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = m


def _add_two_column_line(
    doc: Document,
    style: DocxStyleConfig,
    left_runs: List[Tuple],  # (text, font_size_pt, bold, italic) or + force_primary_font bool
    right_run: Optional[Tuple[str, float, bool, bool]],
    *,
    indent_pt: float = 0,
    space_after_pt: float = 0,
    space_before_pt: float = 0,
    line_spacing_single: bool = False,
) -> None:
    """
    Add a line with left content (multiple styled runs) and right-aligned content.
    Matches PDF flex layout: school+GPA left, dates right; degree left, location right; etc.
    left_runs: [(text, font_size_pt, bold, italic), ...]
    right_run: (text, font_size_pt, bold, italic) or None
    """
    has_left = any(r[0] for r in left_runs)
    has_right = right_run and right_run[0]
    if not has_left and not has_right:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(indent_pt)
    # Always set (including 0); `if space_after_pt` skipped 0 and Word kept default spacing
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    if line_spacing_single:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if has_right:
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(style.two_column_tab_in),
            WD_TAB_ALIGNMENT.RIGHT,
        )

    for run_spec in left_runs:
        if len(run_spec) == 5:
            text, sz, bold, italic, force_primary = run_spec
        else:
            text, sz, bold, italic = run_spec
            force_primary = False
        if not text:
            continue
        run = p.add_run(str(text))
        run.font.size = Pt(sz)
        if force_primary:
            run.font.name = style.font_primary
        else:
            run.font.name = style.font_primary if bold else style.font_secondary
        run.bold = bold
        run.italic = italic

    if has_right and right_run:
        text, sz, bold, italic = right_run
        if text and str(text).strip():
            p.add_run("\t")
            run = p.add_run(str(text).strip())
            run.font.size = Pt(sz)
            run.font.name = style.font_primary if bold else style.font_secondary
            run.bold = bold
            run.italic = italic


def _format_date_range(start_raw, end_raw: str, current: bool) -> str:
    """Format date range for display. Matches builders logic."""
    start_val = str(start_raw) if start_raw else ""
    end_val = str(end_raw) if end_raw else ""
    start_date = format_date_month_year(start_val) if start_val else ""
    end_date = format_date_month_year(end_val) if end_val else ""

    if current and start_date:
        return f"{start_date} - Present"
    if current and not start_date:
        return "Present"
    if start_date and end_date:
        return f"{start_date} - {end_date}"
    if start_date:
        return start_date
    if end_date:
        return end_date
    return ""


def _apply_section_title_bottom_border(p) -> None:
    """
    Draw the divider on the title paragraph itself (no empty paragraph below).
    Avoids an extra 'blank line' in Word that deletes with the rule and feels disconnected.
    """
    p_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "2")  # ~0.25pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_border.append(bottom)
    p._p.get_or_add_pPr().append(p_border)


def _add_header(
    document: Document,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    style_preferences: dict | None = None,
) -> None:
    """Add name and contact line with proper styling."""
    header = resume_data.get("header", {})
    name = f"{header.get('first_name', '')} {header.get('last_name', '')}".strip()
    if name:
        p = document.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(style.name_font_size_pt)
        run.font.name = style.font_primary
        _set_run_character_spacing(run, style.name_letter_spacing_pt or 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(style.name_space_after_pt)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    has_tagline = False
    vis = header.get("visibility", {})
    if vis.get("showTagline", True):
        tagline = (header.get("tagline") or "").strip()
        if tagline:
            has_tagline = True
            p = document.add_paragraph()
            for text, t_bold, t_italic, t_underline in parse_tagline_runs(tagline):
                run = p.add_run(text)
                run.font.size = Pt(style.tagline_font_size_pt)
                run.font.name = style.font_primary
                run.bold = t_bold
                run.italic = t_italic
                run.font.underline = WD_UNDERLINE.SINGLE if t_underline else WD_UNDERLINE.NONE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(style.tagline_space_after_pt)
            _set_line_spacing_multiple(p, style.tagline_line_height)

    # Contact line
    contact_order = header.get("contactOrder", ["email", "phone", "location", "linkedin", "github", "portfolio"])
    field_map = {
        "email": header.get("email", ""),
        "phone": header.get("phone", ""),
        "location": header.get("location", ""),
        "linkedin": header.get("linkedin", ""),
        "github": header.get("github", ""),
        "portfolio": header.get("portfolio", ""),
    }
    visibility = header.get("visibility", {})
    visibility_map = {
        "email": "showEmail", "phone": "showPhone", "location": "showLocation",
        "linkedin": "showLinkedin", "github": "showGithub", "portfolio": "showPortfolio",
    }
    url_disp = resolve_contact_url_display(style_preferences)
    fields = []
    for field_key in contact_order:
        if field_key in field_map:
            val = field_map[field_key]
            vis_key = visibility_map.get(field_key)
            is_visible = visibility.get(vis_key, True) if vis_key else True
            if val and str(val).strip() and is_visible:
                fields.append(
                    format_contact_field_display(
                        field_key,
                        str(val).strip(),
                        contact_url_display=url_disp,
                    )
                )
    if fields:
        p = document.add_paragraph()
        run = p.add_run(" | ".join(fields))
        run.font.size = Pt(style.contact_font_size_pt)
        run.font.name = style.font_primary  # Georgia (contact line), not Times New Roman
        _set_run_character_spacing(run, style.contact_span_letter_spacing_pt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        space_before = (
            style.contact_space_before_after_tagline_pt
            if has_tagline
            else style.contact_space_before_pt
        )
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(style.contact_space_after_pt)
        _set_line_spacing_multiple(p, style.contact_line_height)


def _add_section_title(document: Document, title: str, style: DocxStyleConfig) -> None:
    """Add a section heading with bottom border on the same paragraph; space_after clears content below."""
    p = document.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(style.section_title_font_size_pt)
    run.font.name = style.font_primary
    p.paragraph_format.space_before = Pt(style.section_title_space_before_pt)
    p.paragraph_format.space_after = Pt(style.section_title_space_after_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _apply_section_title_bottom_border(p)


def _add_para(
    document: Document,
    text: str,
    style: DocxStyleConfig,
    *,
    font_size_pt: Optional[float] = None,
    font_name: Optional[str] = None,
    bold: bool = False,
    italic: bool = False,
    indent_pt: float = 0,
    first_line_indent_pt: float = 0,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
    alignment: Optional[WD_ALIGN_PARAGRAPH] = None,
) -> None:
    """Add a styled paragraph."""
    if not text or not str(text).strip():
        return
    p = document.add_paragraph()
    run = p.add_run(str(text).strip())
    run.font.size = Pt(font_size_pt or style.description_font_size_pt)
    if font_name:
        run.font.name = font_name
    else:
        run.font.name = style.font_primary if bold else style.font_secondary
    run.bold = bold
    run.italic = italic
    if indent_pt:
        p.paragraph_format.left_indent = Pt(indent_pt)
    if first_line_indent_pt:
        p.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    _set_line_spacing_multiple(p, style.prose_line_height)
    if alignment is not None:
        p.alignment = alignment


def _parse_description_items(description: str) -> List[Tuple[str, str]]:
    """Parse description into bullet/paragraph items. Matches builders.format_description logic."""
    if not description or not str(description).strip():
        return []
    lines = str(description).split("\n")
    items = []
    non_bullet_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("•"):
            if non_bullet_lines:
                items.append(("paragraph", "\n".join(non_bullet_lines)))
                non_bullet_lines = []
            bullet_text = re.sub(r"^•\s+", "", stripped)
            if bullet_text:
                items.append(("bullet", bullet_text))
        else:
            non_bullet_lines.append(stripped)
    if non_bullet_lines:
        items.append(("paragraph", "\n".join(non_bullet_lines)))
    return items


def _add_description_block(
    doc: Document,
    style: DocxStyleConfig,
    description: str,
    indent_pt: float,
) -> None:
    """
    Add description with bullets set off to the side (matches PDF).
    Word bullets: single line spacing; space before/after from word_bullet_* tokens (see PDF_WORD_SPACING.md).
    Non-bullet lines keep prose_line_height and description_paragraph_space_pt.
    """
    items = _parse_description_items(description)
    if not items:
        return
    bullet_indent = indent_pt + style.description_bullet_indent_pt
    first = True
    for item_type, content in items:
        if not content or not str(content).strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if first:
            if item_type == "bullet":
                p.paragraph_format.space_before = Pt(
                    style.description_block_space_before_pt + style.word_bullet_space_before_pt
                )
            else:
                p.paragraph_format.space_before = Pt(style.description_block_space_before_pt)
            first = False
        else:
            if item_type == "bullet":
                p.paragraph_format.space_before = Pt(style.word_bullet_space_before_pt)
            else:
                p.paragraph_format.space_before = Pt(0)
        if item_type == "bullet":
            # Tab stop at hang_pt (from first-line edge) aligns text; wrapped lines align with left_indent
            p.paragraph_format.left_indent = Pt(bullet_indent)
            p.paragraph_format.first_line_indent = Pt(-style.description_bullet_hang_pt)
            p.paragraph_format.tab_stops.add_tab_stop(
                Pt(style.description_bullet_hang_pt),
                WD_TAB_ALIGNMENT.LEFT,
            )
            p.paragraph_format.space_after = Pt(style.word_bullet_space_after_pt)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run("•\t" + str(content).strip())
        else:
            p.paragraph_format.left_indent = Pt(indent_pt)
            p.paragraph_format.space_after = Pt(style.description_paragraph_space_pt)
            run = p.add_run(str(content).strip())
            _set_line_spacing_multiple(p, style.prose_line_height)
        run.font.size = Pt(style.description_font_size_pt)
        run.font.name = style.font_primary  # Georgia for bullets & description (Word)


# Matches pipeline.py / PDF: header is fixed; body sections follow resume_data.sectionOrder
DEFAULT_DOCX_SECTION_ORDER = ["summary", "education", "experience", "projects", "skills"]
_DOCX_SECTION_KEYS = frozenset(DEFAULT_DOCX_SECTION_ORDER)


def _normalize_docx_section_order(order: Any) -> List[str]:
    """Skip 'header'; dedupe; fall back to default if empty or invalid."""
    if not order or not isinstance(order, list):
        return list(DEFAULT_DOCX_SECTION_ORDER)
    seen = set()
    out: List[str] = []
    for key in order:
        if key == "header":
            continue
        if key in _DOCX_SECTION_KEYS and key not in seen:
            out.append(key)
            seen.add(key)
    if "summary" in out:
        out = ["summary"] + [k for k in out if k != "summary"]
    return out if out else list(DEFAULT_DOCX_SECTION_ORDER)


def _render_docx_summary_section(
    doc: Document,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:
    summary = resume_data.get("summary")
    if not summary or not isinstance(summary, dict):
        return
    text = summary.get("summary", "").strip()
    if not text:
        return
    _add_section_title(doc, section_labels.get("summary", defaults["summary"]), style)
    summary_indent = indent + style.summary_text_padding_left_pt
    _add_para(
        doc, text, style,
        font_size_pt=style.summary_font_size_pt,
        font_name=style.font_primary,
        indent_pt=summary_indent,
        first_line_indent_pt=style.summary_first_line_indent_pt,
        space_after_pt=style.summary_space_after_pt,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _render_docx_education_section(
    doc: Document,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:
    education = resume_data.get("education") or []
    if not education:
        return
    _add_section_title(doc, section_labels.get("education", defaults["education"]), style)
    for edu in education:
        degree = edu.get("degree", "")
        discipline = edu.get("discipline", "")
        minor = edu.get("minor", "")
        degree_text = f"{degree} in {discipline}" if discipline else degree
        if minor:
            degree_text += f", Minor in {minor}"

        start_raw = edu.get("start_date") or edu.get("startDate") or ""
        end_raw = edu.get("end_date") or edu.get("endDate") or ""
        date_range = _format_date_range(start_raw, end_raw, edu.get("current", False))
        gpa = edu.get("gpa", "")
        gpa_text = f" (GPA: {gpa})" if gpa else ""

        left_runs = [
            (edu.get("school", ""), style.school_name_font_size_pt, True, False),
        ]
        if gpa_text:
            left_runs.append(
                (gpa_text, style.school_gpa_font_size_pt, False, True, True)
            )
        _add_two_column_line(
            doc, style,
            left_runs,
            (date_range, style.school_meta_font_size_pt, False, True) if date_range else None,
            indent_pt=indent,
            space_after_pt=style.school_name_line_space_after_pt,
        )

        loc = edu.get("location", "")
        deg_sz = style.education_degree_line_font_size_pt
        _add_two_column_line(
            doc, style,
            [(degree_text, deg_sz, False, True)],
            (loc, style.school_meta_font_size_pt, False, True) if loc else None,
            indent_pt=indent,
            space_after_pt=style.school_line_space_after_pt,
        )
        for sub_title, sub_content in (edu.get("subsections") or {}).items():
            if sub_content and str(sub_content).strip():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(indent)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(style.word_highlight_space_after_pt)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if sub_title:
                    r1 = p.add_run(f"{sub_title}: ")
                    r1.bold = True
                    r1.font.size = Pt(style.highlight_font_size_pt)
                    r1.font.name = style.font_primary
                r2 = p.add_run(str(sub_content).strip())
                r2.font.size = Pt(style.highlight_font_size_pt)
                r2.font.name = style.font_primary


def _render_docx_experience_section(
    doc: Document,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:
    experience = resume_data.get("experience") or []
    if not experience:
        return
    _add_section_title(doc, section_labels.get("experience", defaults["experience"]), style)
    for exp in experience:
        start_raw = exp.get("start_date") or exp.get("startDate") or ""
        end_raw = exp.get("end_date") or exp.get("endDate") or ""
        date_range = _format_date_range(start_raw, end_raw, exp.get("current", False))

        _add_two_column_line(
            doc, style,
            [(exp.get("title", ""), style.experience_title_font_size_pt, True, False)],
            (date_range, style.experience_meta_font_size_pt, False, True) if date_range else None,
            indent_pt=indent,
            space_before_pt=style.experience_line_space_before_pt,
            space_after_pt=style.experience_line_space_pt,
            line_spacing_single=True,
        )

        company = exp.get("company", "")
        skills = exp.get("skills", "")
        loc = exp.get("location", "")
        left_runs = [(company, style.experience_meta_font_size_pt, True, False)]
        if skills:
            left_runs.append((f" | {skills}", style.experience_meta_font_size_pt, False, True))
        _add_two_column_line(
            doc, style,
            left_runs,
            (loc, style.experience_meta_font_size_pt, False, True) if loc else None,
            indent_pt=indent,
            space_after_pt=style.company_line_space_after_pt,
            line_spacing_single=True,
        )
        desc = exp.get("description", "")
        if desc and str(desc).strip():
            _add_description_block(doc, style, str(desc), indent)


def _render_docx_projects_section(
    doc: Document,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:
    projects = resume_data.get("projects") or []
    if not projects:
        return
    _add_section_title(doc, section_labels.get("projects", defaults["projects"]), style)
    for proj in projects:
        title = proj.get("title", "")
        tech = proj.get("tech_stack") or proj.get("techStack") or []
        tech_str = ", ".join(tech) if isinstance(tech, list) else str(tech or "")
        url = proj.get("url", "")

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(indent)
        p.paragraph_format.space_before = Pt(style.project_title_space_before_pt)
        p.paragraph_format.space_after = Pt(style.project_line_space_pt)
        _set_line_spacing_multiple(p, 1.0)
        r1 = p.add_run(title)
        r1.bold = True
        r1.font.size = Pt(style.project_title_font_size_pt)
        r1.font.name = style.font_primary
        if tech_str or url:
            extras = []
            if tech_str:
                extras.append(tech_str)
            if url:
                extras.append(url)
            p.add_run(" | ")
            r2 = p.add_run(" | ".join(extras))
            r2.italic = True
            r2.font.size = Pt(style.project_title_font_size_pt)
            r2.font.name = style.font_secondary
        desc = proj.get("description", "")
        if desc and str(desc).strip():
            _add_description_block(doc, style, str(desc), indent)


def _render_docx_skills_section(
    doc: Document,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:
    skills = resume_data.get("skills") or []
    skills_category_order = resume_data.get("skillsCategoryOrder") or []
    if not skills:
        return

    skills_by_cat: Dict[str, List[str]] = {}
    uncategorized: List[str] = []
    for s in skills:
        if isinstance(s, dict):
            cat = s.get("category") or s.get("Category")
            name = s.get("name") or s.get("Name", "")
        else:
            cat = None
            name = str(s)
        if not name:
            continue
        if cat:
            skills_by_cat.setdefault(cat, []).append(name)
        else:
            uncategorized.append(name)

    _add_section_title(doc, section_labels.get("skills", defaults["skills"]), style)
    order = [c for c in skills_category_order if c in skills_by_cat]
    for c in sorted(skills_by_cat.keys()):
        if c not in order:
            order.append(c)
    for cat in order:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(indent)
        p.paragraph_format.space_after = Pt(style.skill_line_space_pt)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r1 = p.add_run(f"{cat}: ")
        r1.bold = True
        r1.font.size = Pt(style.skill_category_font_size_pt)
        r1.font.name = style.font_primary
        r2 = p.add_run(", ".join(skills_by_cat[cat]))
        r2.font.size = Pt(style.skill_names_font_size_pt)
        r2.font.name = style.font_primary
    if uncategorized:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(indent)
        p.paragraph_format.space_after = Pt(style.skill_line_space_pt)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(", ".join(uncategorized))
        r.font.size = Pt(style.skill_names_font_size_pt)
        r.font.name = style.font_primary


def build_docx(
    resume_data: Dict[str, Any],
    template_name: str = "classic",
    style_preferences: dict | None = None,
) -> bytes:
    """
    Build a styled Word document from resume_data.
    Uses docx_styles for template-consistent formatting (margins, fonts, spacing).
    Body sections follow resume_data.sectionOrder (same as PDF/HTML), after the header.
    """
    style = get_styles(template_name, style_preferences)
    doc = Document()

    # Page margins (match .resume padding)
    sect = doc.sections[0]
    sect.top_margin = Inches(style.margin_top_in)
    sect.bottom_margin = Inches(style.margin_bottom_in)
    sect.left_margin = Inches(style.margin_left_in)
    sect.right_margin = Inches(style.margin_right_in)

    # Header (always first)
    _add_header(doc, resume_data, style, style_preferences)

    section_labels = resume_data.get("sectionLabels", {})
    defaults = {
        "summary": "Professional Summary",
        "education": "Education",
        "experience": "Experience",
        "projects": "Projects",
        "skills": "Skills",
    }

    indent = style.section_content_indent_pt

    section_order = _normalize_docx_section_order(resume_data.get("sectionOrder"))
    renderers = {
        "summary": _render_docx_summary_section,
        "education": _render_docx_education_section,
        "experience": _render_docx_experience_section,
        "projects": _render_docx_projects_section,
        "skills": _render_docx_skills_section,
    }
    for section_key in section_order:
        fn = renderers.get(section_key)
        if fn:
            fn(doc, resume_data, style, indent, section_labels, defaults)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
