# Single-column (and shared) body sections: Summary, Education, Experience, Projects, Skills.

from __future__ import annotations

from typing import Any, Dict, List

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt

from ..shared.dates import format_date_range
from ..shared.skills import skills_group_ordered
from .docx_description import _add_description_block
from .docx_layout import (
    _add_para,
    _add_two_column_line,
    _apply_section_title_bottom_border,
)
from .docx_run_style import _apply_run_resume_color, _set_line_spacing_multiple
from .docx_styles import DocxStyleConfig


# --- Handle Section Title ---
def _add_section_title(doc: Any, title: str, style: DocxStyleConfig) -> None:
    # Add a section heading with bottom border on the same paragraph; space_after clears content below.
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(style.section_title_font_size_pt)
    run.font.name = style.font_primary
    _apply_run_resume_color(run, style, "resume_text_color_emphasis")
    p.paragraph_format.space_before = Pt(style.section_title_space_before_pt)
    p.paragraph_format.space_after = Pt(style.section_title_space_after_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _apply_section_title_bottom_border(p)


# --- Handle Normalize Section Order ---
# Matches pipeline.py / PDF: header is fixed; body sections follow resume_data.sectionOrder
DEFAULT_DOCX_SECTION_ORDER = ["summary", "education", "experience", "projects", "skills"]
_DOCX_SECTION_KEYS = frozenset(DEFAULT_DOCX_SECTION_ORDER)


def _normalize_docx_section_order(order: Any) -> List[str]:
    # Skip 'header'; dedupe; fall back to default if empty or invalid.
    if not order or not isinstance(order, list):
        return list(DEFAULT_DOCX_SECTION_ORDER)
    seen = set()
    out: List[str] = []
    for key in order:
        if key == "header":
            continue
        if key in _DOCX_SECTION_KEYS and key not in seen:
            out.append(key)
            seen.add(key)
    if "summary" in out:
        out = ["summary"] + [k for k in out if k != "summary"]
    return out if out else list(DEFAULT_DOCX_SECTION_ORDER)


# --- Handle Render Summary Section ---
def _render_docx_summary_section(
    doc: Any,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:
    # Get summary.
    summary = resume_data.get("summary")
    if not summary or not isinstance(summary, dict):
        return
    text = summary.get("summary", "").strip()
    if not text:
        return

    # Build Summary Section w/ Title and Content.
    _add_section_title(doc, section_labels.get("summary", defaults["summary"]), style)
    summary_indent = indent + style.summary_text_padding_left_pt
    _add_para(
        doc, text, style,
        font_size_pt=style.summary_font_size_pt,
        font_name=style.font_primary,
        indent_pt=summary_indent,
        first_line_indent_pt=style.summary_first_line_indent_pt,
        space_after_pt=style.summary_space_after_pt,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_height_multiple=style.summary_line_height,
    )


# --- Handle Render Education Section ---
def _render_docx_education_section(
    doc: Any,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:
    # Get education.
    education = resume_data.get("education") or []
    if not education:
        return
    _add_section_title(doc, section_labels.get("education", defaults["education"]), style)
    
    # Iterate over each education entry.
    for edu in education:
        # Get Education Details.
        degree = edu.get("degree", "")
        discipline = edu.get("discipline", "")
        minor = edu.get("minor", "")
        degree_text = f"{degree} in {discipline}" if discipline else degree
        if minor:
            degree_text += f", Minor in {minor}"

        start_raw = edu.get("start_date") or edu.get("startDate") or ""
        end_raw = edu.get("end_date") or edu.get("endDate") or ""
        date_range = format_date_range(start_raw, end_raw, edu.get("current", False))
        gpa = edu.get("gpa", "")
        gpa_text = f" (GPA: {gpa})" if gpa else ""

        left_runs = [
            (edu.get("school", ""), style.school_name_font_size_pt, True, False),
        ]
        if gpa_text:
            left_runs.append(
                (gpa_text, style.school_gpa_font_size_pt, False, True, False)
            )
        _add_two_column_line(
            doc, style,
            left_runs,
            (date_range, style.school_meta_font_size_pt, False, True) if date_range else None,
            indent_pt=indent,
            space_after_pt=style.school_name_line_space_after_pt,
        )

        loc = edu.get("location", "")
        deg_sz = style.education_degree_line_font_size_pt
        _add_two_column_line(
            doc, style,
            [(degree_text, deg_sz, False, True)],
            (loc, style.school_meta_font_size_pt, False, True) if loc else None,
            indent_pt=indent,
            space_after_pt=style.school_line_space_after_pt,
        )
        hl_between = float(getattr(style, "highlights_gap_pt", 4.5) or 4.5)
        hl_ord = 0

        # Build Highlights Lines.
        for sub_title, sub_content in (edu.get("subsections") or {}).items():
            if sub_content and str(sub_content).strip():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(indent)
                p.paragraph_format.space_before = Pt(0 if hl_ord == 0 else hl_between)
                hl_ord += 1
                p.paragraph_format.space_after = Pt(style.word_highlight_space_after_pt)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                if sub_title:
                    r1 = p.add_run(f"{sub_title}: ")
                    r1.bold = True
                    r1.font.size = Pt(style.highlight_font_size_pt)
                    r1.font.name = style.font_primary
                    _apply_run_resume_color(r1, style, "resume_text_color_highlight_title")
                r2 = p.add_run(str(sub_content).strip())
                r2.font.size = Pt(style.highlight_font_size_pt)
                r2.font.name = style.font_primary
                _apply_run_resume_color(r2, style, "resume_text_color_primary")


# --- Handle Render Experience Section ---
def _render_docx_experience_section(
    doc: Any,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:

    # Get experience.
    experience = resume_data.get("experience") or []
    if not experience:
        return
    _add_section_title(doc, section_labels.get("experience", defaults["experience"]), style)

    # Iterate over each experience entry.
    for exp in experience:
        start_raw = exp.get("start_date") or exp.get("startDate") or ""
        end_raw = exp.get("end_date") or exp.get("endDate") or ""
        date_range = format_date_range(start_raw, end_raw, exp.get("current", False))

        # Build Experience Line.
        _add_two_column_line(
            doc, style,
            [(exp.get("title", ""), style.experience_title_font_size_pt, True, False)],
            (date_range, style.experience_meta_font_size_pt, False, True) if date_range else None,
            indent_pt=indent,
            space_before_pt=style.experience_line_space_before_pt,
            space_after_pt=style.experience_line_space_pt,
            line_spacing_single=True,
        )

        # Build Company Line.
        company = exp.get("company", "")
        skills = exp.get("skills", "")
        loc = exp.get("location", "")
        left_runs = [(company, style.experience_meta_font_size_pt, True, False)]
        if skills:
            left_runs.append((f" | {skills}", style.experience_meta_font_size_pt, False, True))
        _add_two_column_line(
            doc, style,
            left_runs,
            (loc, style.experience_meta_font_size_pt, False, True) if loc else None,
            indent_pt=indent,
            space_after_pt=style.company_line_space_after_pt,
            line_spacing_single=True,
        )

        # Build Description Details.
        desc = exp.get("description", "")
        if desc and str(desc).strip():
            _add_description_block(doc, style, str(desc), indent)


# --- Handle Render Projects Section ---
def _render_docx_projects_section(
    doc: Any,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
    *,
    sidebar_main_column: bool = False,
) -> None:

    # Get projects.
    projects = resume_data.get("projects") or []
    if not projects:
        return

    # Build Projects Section w/ Title.
    _add_section_title(doc, section_labels.get("projects", defaults["projects"]), style)
    meta_sz = float(style.project_title_font_size_pt)
    for proj in projects:
        # Get Project Details.
        title = proj.get("title", "")
        tech = proj.get("tech_stack") or proj.get("techStack") or []
        tech_str = ", ".join(tech) if isinstance(tech, list) else str(tech or "")
        url = str(proj.get("url") or "").strip()

        # Build Project Line.
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(indent)
        p.paragraph_format.space_before = Pt(style.project_title_space_before_pt)
        p.paragraph_format.space_after = (
            Pt(style.project_line_space_pt * 0.35)
            if sidebar_main_column and (tech_str or url)
            else Pt(style.project_line_space_pt)
        )
        _set_line_spacing_multiple(p, 1.0)
        r1 = p.add_run(title)
        r1.bold = True
        r1.font.size = Pt(meta_sz)
        r1.font.name = style.font_primary
        _apply_run_resume_color(r1, style, "resume_text_color_emphasis")

        # Build Tech Stack and URL Line.
        if sidebar_main_column and (tech_str or url):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Pt(indent)
            p2.paragraph_format.space_after = Pt(style.project_line_space_pt)
            _set_line_spacing_multiple(p2, 1.0)
            if url:
                p2.paragraph_format.tab_stops.add_tab_stop(
                    Inches(style.two_column_tab_in),
                    WD_TAB_ALIGNMENT.RIGHT,
                )
            if tech_str:
                r_t = p2.add_run(tech_str)
                r_t.italic = True
                r_t.font.size = Pt(meta_sz)
                r_t.font.name = style.font_secondary
                _apply_run_resume_color(r_t, style, "resume_text_color_secondary")
            if url:
                p2.add_run("\t")
                r_u = p2.add_run(url)
                r_u.italic = True
                r_u.font.size = Pt(meta_sz)
                r_u.font.name = style.font_secondary
                _apply_run_resume_color(r_u, style, "resume_text_color_secondary")
        elif tech_str or url:
            extras = []
            if tech_str:
                extras.append(tech_str)
            if url:
                extras.append(url)
            p.add_run(" | ")
            r2 = p.add_run(" | ".join(extras))
            r2.italic = True
            r2.font.size = Pt(meta_sz)
            r2.font.name = style.font_secondary
            _apply_run_resume_color(r2, style, "resume_text_color_secondary")

        # Build Description Details.
        desc = proj.get("description", "")
        if desc and str(desc).strip():
            _add_description_block(doc, style, str(desc), indent)


# --- Handle Render Skills Section ---
def _render_docx_skills_section(
    doc: Any,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:

    # Get skills.
    skills = resume_data.get("skills") or []
    skills_category_order = resume_data.get("skillsCategoryOrder") or []
    if not skills:
        return

    # Build Skills by Category.
    skills_by_cat: Dict[str, List[str]] = {}
    uncategorized: List[str] = []
    for s in skills:
        if isinstance(s, dict):
            cat = s.get("category") or s.get("Category")
            name = s.get("name") or s.get("Name", "")
        else:
            cat = None
            name = str(s)
        if not name:
            continue
        if cat:
            skills_by_cat.setdefault(cat, []).append(name)
        else:
            uncategorized.append(name)

    # Build Skills Section w/ Title.
    _add_section_title(doc, section_labels.get("skills", defaults["skills"]), style)
    order = [c for c in skills_category_order if c in skills_by_cat]
    for c in sorted(skills_by_cat.keys()):
        if c not in order:
            order.append(c)

    # Iterate over each skill category.
    for cat in order:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(indent)
        p.paragraph_format.space_after = Pt(style.skill_line_space_pt)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r1 = p.add_run(f"{cat}: ")
        r1.bold = True
        r1.font.size = Pt(style.skill_category_font_size_pt)
        r1.font.name = style.font_primary
        _apply_run_resume_color(r1, style, "resume_text_color_emphasis")
        r2 = p.add_run(", ".join(skills_by_cat[cat]))
        r2.font.size = Pt(style.skill_names_font_size_pt)
        r2.font.name = style.font_primary
        _apply_run_resume_color(r2, style, "resume_text_color_primary")

    # Build Uncategorized Skills.
    if uncategorized:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(indent)
        p.paragraph_format.space_after = Pt(style.skill_line_space_pt)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(", ".join(uncategorized))
        r.font.size = Pt(style.skill_names_font_size_pt)
        r.font.name = style.font_primary
        _apply_run_resume_color(r, style, "resume_text_color_primary")
