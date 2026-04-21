# Paragraph Layout Primitives: Two-column lines, section title rule, generic prose paragraphs.

from __future__ import annotations

from typing import Any, List, Optional, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .docx_run_style import _apply_run_resume_color, _set_line_spacing_multiple
from .docx_styles import DocxStyleConfig


# --- Handle Two-column Line ---
def _add_two_column_line(
    doc: Any,
    style: DocxStyleConfig,
    left_runs: List[Tuple],  # (text, font_size_pt, bold, italic) or + force_primary_font bool
    right_run: Optional[Tuple[str, float, bool, bool]],
    *,
    indent_pt: float = 0,
    space_after_pt: float = 0,
    space_before_pt: float = 0,
    line_spacing_single: bool = False,
) -> None:

    # Check if there is any left or right content.
    has_left = any(r[0] for r in left_runs)
    has_right = right_run and right_run[0]
    if not has_left and not has_right:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(indent_pt)
    # Always set (including 0); if space_after_pt skipped 0 and Word kept default spacing.
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    if line_spacing_single:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add tab stop for right-aligned content.
    if has_right:
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(style.two_column_tab_in),
            WD_TAB_ALIGNMENT.RIGHT,
        )

    # Iterate over each left run.
    for run_spec in left_runs:
        if len(run_spec) == 5:
            text, sz, bold, italic, force_primary = run_spec
        else:
            text, sz, bold, italic = run_spec
            force_primary = False
        if not text:
            continue
        run = p.add_run(str(text))
        run.font.size = Pt(sz)
        if force_primary:
            run.font.name = style.font_primary
        else:
            run.font.name = style.font_primary if bold else style.font_secondary
        run.bold = bold
        run.italic = italic
        if bold:
            _apply_run_resume_color(run, style, "resume_text_color_emphasis")
        elif italic:
            _apply_run_resume_color(run, style, "resume_text_color_secondary")
        else:
            _apply_run_resume_color(run, style, "resume_text_color_primary")

    # Add right run (if visible).
    if has_right and right_run:
        text, sz, bold, italic = right_run
        if text and str(text).strip():
            p.add_run("\t")
            run = p.add_run(str(text).strip())
            run.font.size = Pt(sz)
            run.font.name = style.font_primary if bold else style.font_secondary
            run.bold = bold
            run.italic = italic
            _apply_run_resume_color(run, style, "resume_text_color_secondary")


# --- Handle Section Title Bottom Border ---
def _apply_section_title_bottom_border(p) -> None:
    # Draw the divider on the title paragraph itself (no empty paragraph below).
    # Avoids an extra 'blank line' in Word that deletes with the rule and feels disconnected.

    p_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "2")  # ~0.25pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_border.append(bottom)
    p._p.get_or_add_pPr().append(p_border)


# --- Handle Generic Prose Paragraph ---
def _add_para(
    doc: Any,
    text: str,
    style: DocxStyleConfig,
    *,
    font_size_pt: Optional[float] = None,
    font_name: Optional[str] = None,
    bold: bool = False,
    italic: bool = False,
    indent_pt: float = 0,
    first_line_indent_pt: float = 0,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
    alignment: Optional[WD_ALIGN_PARAGRAPH] = None,
    line_height_multiple: Optional[float] = None,
) -> None:

    # Check if there is any text.
    if not text or not str(text).strip():
        return
    p = doc.add_paragraph()

    # Add run --- 
    run = p.add_run(str(text).strip())
    run.font.size = Pt(font_size_pt or style.description_font_size_pt)
    if font_name:
        run.font.name = font_name
    else:
        run.font.name = style.font_primary if bold else style.font_secondary
    run.bold = bold
    run.italic = italic
    _apply_run_resume_color(run, style, "resume_text_color_primary")
    if indent_pt:
        p.paragraph_format.left_indent = Pt(indent_pt)
    if first_line_indent_pt:
        p.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    _set_line_spacing_multiple(
        p, line_height_multiple if line_height_multiple is not None else style.prose_line_height
    )
    if alignment is not None:
        p.alignment = alignment

    # -----------
