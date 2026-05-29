# Timeline-split DOCX layout.
#
# Word is table-first, not CSS-grid-first. This builder intentionally uses one
# fixed-width body table with four columns:
#   1. left rail: contact + skills
#   2. marker: section icon images
#   3. rule: vertical line
#   4. main: profile, work, projects, education
#
# The marker and rule are intentionally separate. Word table borders and inline
# images do not overlap reliably, so this keeps the spine visually consistent
# without depending on floating-position XML.

from __future__ import annotations

from dataclasses import replace
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from ..html import format_contact_field_display, resolve_contact_url_display
from ..html.header import iter_visible_contact_fields
from ..icons.docx_rail_png import contact_rail_icon_png_bytes
from ..layouts.timeline_split import timeline_main_column_order
from ..shared.skills import skills_group_ordered
from ..shared.tagline import parse_tagline_runs
from ..shared.template_slug import resolve_template_folder
from .docx_layout import _apply_section_title_bottom_border
from .docx_run_style import (
    _apply_run_resume_color,
    _hex_rgb_6_for_word,
    _set_line_spacing_multiple,
    _set_run_character_spacing,
)
from .docx_sections import (
    _render_docx_education_section,
    _render_docx_experience_section,
    _render_docx_projects_section,
    _render_docx_summary_section,
)
from .docx_sidebar_split import _remove_singleton_empty_paragraph, _set_cell_margins_inches
from .docx_styles import DocxStyleConfig


_TIMELINE_ICON_NAMES = {
    "summary": ("summary.png", "profile.png"),
    "experience": ("experience.png", "work.png"),
    "projects": ("projects.png", "project.png"),
    "education": ("education.png", "school.png"),
}


def _remove_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        existing = borders.find(qn(f"w:{side}"))
        if existing is None:
            existing = OxmlElement(f"w:{side}")
            borders.append(existing)
        existing.set(qn("w:val"), "nil")


def _set_cell_width(cell: Any, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.insert(0, tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(round(float(width_in) * 1440))))


def _set_table_geometry(table: Any, widths_in: Iterable[float]) -> None:
    widths = [max(0.1, float(width)) for width in widths_in]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    try:
        table.allow_autofit = False
    except AttributeError:
        pass

    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(round(sum(widths) * 1440))))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for old in list(grid.findall(qn("w:gridCol"))):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                _set_cell_width(row.cells[index], width)

    _remove_table_borders(table)


def _set_cell_border(cell: Any, side: str, *, color: str, width_pt: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for old in list(borders.findall(qn(f"w:{side}"))):
        borders.remove(old)

    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(max(2, min(96, int(round(width_pt * 8))))))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), _hex_rgb_6_for_word(color) or "1F2937")
    borders.append(border)


def _section_has_content(section_key: str, resume_data: Dict[str, Any]) -> bool:
    if section_key == "summary":
        summary = resume_data.get("summary")
        return isinstance(summary, dict) and bool(str(summary.get("summary") or "").strip())
    value = resume_data.get(section_key)
    return isinstance(value, list) and len(value) > 0


def _timeline_main_keys(resume_data: Dict[str, Any]) -> List[str]:
    return [
        key
        for key in timeline_main_column_order(resume_data)
        if _section_has_content(key, resume_data)
    ]


def _timeline_section_labels(resume_data: Dict[str, Any]) -> Dict[str, str]:
    raw = resume_data.get("sectionLabels") if isinstance(resume_data.get("sectionLabels"), dict) else {}
    labels = dict(raw)
    if not str(labels.get("summary") or "").strip() or labels.get("summary") == "Professional Summary":
        labels["summary"] = "Profile"
    if not str(labels.get("experience") or "").strip() or labels.get("experience") == "Experience":
        labels["experience"] = "Work Experience"
    return labels


def _add_timeline_header(document: Document, resume_data: Dict[str, Any], style: DocxStyleConfig) -> None:
    header = resume_data.get("header", {})
    name = f"{header.get('first_name', '')} {header.get('last_name', '')}".strip()

    if name:
        p = document.add_paragraph()
        run = p.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(style.name_font_size_pt)
        run.font.name = style.font_primary
        _apply_run_resume_color(run, style, "resume_text_color_emphasis")
        _set_run_character_spacing(run, style.name_letter_spacing_pt or 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(style.name_space_after_pt)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    visibility = header.get("visibility", {})
    tagline = (header.get("tagline") or "").strip()
    if visibility.get("showTagline", True) and tagline:
        p = document.add_paragraph()
        for text, bold, italic, underline in parse_tagline_runs(tagline):
            run = p.add_run(text.upper())
            run.font.size = Pt(style.tagline_font_size_pt)
            run.font.name = style.font_primary
            run.bold = bold
            run.italic = italic
            run.font.underline = WD_UNDERLINE.SINGLE if underline else WD_UNDERLINE.NONE
            _apply_run_resume_color(run, style, "resume_text_color_primary")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(style.tagline_space_after_pt)
        _set_line_spacing_multiple(p, style.tagline_line_height)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(1)
    rule.paragraph_format.space_after = Pt(style.resume_header_margin_bottom_pt)
    _apply_section_title_bottom_border(rule, style)


def _add_left_section_title(cell: Any, title: str, style: DocxStyleConfig, *, first: bool = False) -> None:
    p = cell.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(max(8, style.section_title_font_size_pt - 0.5))
    run.font.name = style.font_primary
    _apply_run_resume_color(run, style, "resume_text_color_emphasis")
    _set_run_character_spacing(run, style.section_title_letter_spacing_pt)
    p.paragraph_format.space_before = Pt(0 if first else style.timeline_left_section_gap_pt)
    p.paragraph_format.space_after = Pt(style.section_title_space_after_pt * 0.75)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _apply_section_title_bottom_border(p, style)


def _render_contact_rail(
    cell: Any,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    style_preferences: dict | None,
    *,
    template_slug: str,
) -> None:
    fields = iter_visible_contact_fields(resume_data.get("header", {}))
    if not fields:
        return

    _add_left_section_title(cell, "Contact", style, first=True)
    url_display = resolve_contact_url_display(style_preferences)

    for field_key, value in fields:
        display = format_contact_field_display(
            field_key,
            value,
            contact_url_display=url_display,
        )
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4.5)
        _set_line_spacing_multiple(p, style.contact_line_height)

        icon_pt = float(getattr(style, "timeline_contact_icon_size_pt", 0) or 0)
        icon_gap = float(getattr(style, "timeline_contact_icon_text_gap_pt", 4.5) or 0)
        png = contact_rail_icon_png_bytes(template_slug, field_key) if icon_pt > 0 else None
        if png:
            try:
                r_icon = p.add_run()
                r_icon.add_picture(BytesIO(png), width=Pt(icon_pt), height=Pt(icon_pt))
                if icon_gap > 0:
                    n_nbsp = max(1, min(6, int(round(icon_gap / 2.5))))
                    r_gap = p.add_run("\u00a0" * n_nbsp)
                    r_gap.font.size = Pt(style.contact_font_size_pt)
                    r_gap.font.name = style.font_primary
            except (OSError, ValueError):
                pass

        r_value = p.add_run(display)
        r_value.font.size = Pt(style.contact_font_size_pt)
        r_value.font.name = style.font_primary
        _apply_run_resume_color(r_value, style, "resume_text_color_primary")


def _render_skills_rail(cell: Any, resume_data: Dict[str, Any], style: DocxStyleConfig) -> None:
    groups = skills_group_ordered(
        resume_data.get("skills") or [],
        resume_data.get("skillsCategoryOrder") or [],
    )
    if not groups:
        return

    _add_left_section_title(cell, "Skills", style)
    for category, names in groups:
        if not names:
            continue
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(style.skill_line_space_pt)
        _set_line_spacing_multiple(p, style.skill_line_height)
        if category:
            r_cat = p.add_run(f"{category}: ")
            r_cat.bold = True
            r_cat.font.size = Pt(style.skill_category_font_size_pt)
            r_cat.font.name = style.font_primary
            _apply_run_resume_color(r_cat, style, "resume_text_color_emphasis")
        r_names = p.add_run(", ".join(names))
        r_names.font.size = Pt(style.skill_names_font_size_pt)
        r_names.font.name = style.font_primary
        _apply_run_resume_color(r_names, style, "resume_text_color_primary")


def _timeline_icon_path(template_slug: str, section_key: str) -> Path | None:
    icon_dir = resolve_template_folder(template_slug) / "docx_icons"
    for filename in _TIMELINE_ICON_NAMES.get(section_key, (f"{section_key}.png",)):
        path = icon_dir / filename
        if path.is_file():
            return path
    return None


def _add_marker(cell: Any, section_key: str, style: DocxStyleConfig, template_slug: str) -> None:
    p = cell.add_paragraph()
    # The rule is the next narrow column. Right-aligning the icon tucks it
    # against that rule while avoiding Word's unreliable image/border overlap.
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(style.section_title_space_before_pt)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    icon_path = _timeline_icon_path(template_slug, section_key)
    if icon_path:
        try:
            run = p.add_run()
            run.add_picture(
                str(icon_path),
                width=Pt(style.timeline_marker_size_pt),
                height=Pt(style.timeline_marker_size_pt),
            )
            return
        except (OSError, ValueError):
            pass

    run = p.add_run("\u25cf")
    run.font.size = Pt(style.timeline_marker_size_pt)
    run.font.name = style.font_primary
    _apply_run_resume_color(run, style, "resume_text_color_emphasis")


def _render_main_section(
    cell: Any,
    section_key: str,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
) -> None:
    indent = style.section_content_indent_pt
    if section_key == "summary":
        _render_docx_summary_section(cell, resume_data, style, indent, section_labels, defaults)
    elif section_key == "experience":
        _render_docx_experience_section(cell, resume_data, style, indent, section_labels, defaults)
    elif section_key == "projects":
        _render_docx_projects_section(
            cell,
            resume_data,
            style,
            indent,
            section_labels,
            defaults,
            sidebar_main_column=True,
        )
    elif section_key == "education":
        _render_docx_education_section(cell, resume_data, style, indent, section_labels, defaults)


def _merge_left_rail(table: Any) -> Any:
    if len(table.rows) == 1:
        return table.rows[0].cells[0]
    return table.rows[0].cells[0].merge(table.rows[-1].cells[0])


def build_docx_timeline_split_document(
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    style_preferences: dict | None = None,
    *,
    template_slug: str = "timeline-split",
    docx_max_pages: int | None = None,
) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(style.margin_top_in)
    section.bottom_margin = Inches(style.margin_bottom_in)
    section.left_margin = Inches(style.margin_left_in)
    section.right_margin = Inches(style.margin_right_in)

    _add_timeline_header(doc, resume_data, style)

    main_keys = _timeline_main_keys(resume_data)
    rows = max(1, len(main_keys))
    usable_in = (
        float(section.page_width.inches)
        - float(section.left_margin.inches)
        - float(section.right_margin.inches)
    )
    left_in = float(style.timeline_left_width_in)
    marker_in = max(0.3, float(style.timeline_marker_size_pt) / 72.0 + 0.04)
    rule_in = max(0.035, min(0.08, float(style.timeline_line_width_pt) / 72.0 + 0.035))
    spine_in = max(float(style.timeline_spine_width_in), marker_in + rule_in)
    content_in = max(1.5, usable_in - left_in - spine_in)

    table = doc.add_table(rows=rows, cols=4)
    _set_table_geometry(table, [left_in, marker_in, rule_in, content_in])

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    left_cell = _merge_left_rail(table)
    _remove_singleton_empty_paragraph(left_cell)
    _set_cell_margins_inches(left_cell, top=0, left=0, bottom=0, right=0.16)
    _render_contact_rail(
        left_cell,
        resume_data,
        style,
        style_preferences,
        template_slug=template_slug,
    )
    _render_skills_rail(left_cell, resume_data, style)

    section_labels = _timeline_section_labels(resume_data)
    defaults = {
        "summary": "Profile",
        "education": "Education",
        "experience": "Work Experience",
        "projects": "Projects",
        "skills": "Skills",
    }
    body_gap_in = max(0.0, float(style.timeline_body_column_gap_pt) / 72.0)
    style_main = replace(
        style,
        two_column_tab_in=max(1.25, content_in - body_gap_in - 0.03),
    )

    for row_index, section_key in enumerate(main_keys):
        marker_cell = table.rows[row_index].cells[1]
        rule_cell = table.rows[row_index].cells[2]
        content_cell = table.rows[row_index].cells[3]
        _remove_singleton_empty_paragraph(marker_cell)
        _remove_singleton_empty_paragraph(rule_cell)
        _remove_singleton_empty_paragraph(content_cell)
        _set_cell_margins_inches(marker_cell, top=0, left=0, bottom=0, right=0)
        _set_cell_margins_inches(rule_cell, top=0, left=0, bottom=0, right=0)
        _set_cell_margins_inches(content_cell, top=0, left=body_gap_in, bottom=0, right=0)
        _set_cell_border(
            rule_cell,
            "left",
            color=style.timeline_line_color,
            width_pt=style.timeline_line_width_pt,
        )
        _add_marker(marker_cell, section_key, style, template_slug)
        _render_main_section(
            content_cell,
            section_key,
            resume_data,
            style_main,
            section_labels,
            defaults,
        )

        if row_index + 1 < len(main_keys) and content_cell.paragraphs:
            content_cell.paragraphs[-1].paragraph_format.space_after = Pt(
                style.timeline_section_gap_pt
            )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
