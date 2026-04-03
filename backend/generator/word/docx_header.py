# Document Header(s): Single-column centered block; Sidebar-split rail header.

from __future__ import annotations

from io import BytesIO
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
from docx.shared import Pt

from ..html import format_contact_field_display, resolve_contact_url_display
from ..icons.docx_rail_png import contact_rail_icon_png_bytes
from ..shared.tagline import parse_tagline_runs
from .docx_run_style import _apply_run_resume_color, _set_line_spacing_multiple, _set_run_character_spacing
from .docx_styles import DocxStyleConfig


# --- Handle Sidebar Rail Header ---
def _add_sidebar_rail_header(
    cell: Any,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    style_preferences: dict | None = None,
    *,
    template_slug: str = "sidebar",
) -> None:

    # Calculate rail padding.
    rail_pad_pt = float(getattr(style, "sidebar_pad_x_in", 0.0) or 0.0) * 72.0
    header = resume_data.get("header", {})
    name = f"{header.get('first_name', '')} {header.get('last_name', '')}".strip()

    # Add name.
    if name:
        p = cell.add_paragraph()
        if rail_pad_pt:
            p.paragraph_format.left_indent = Pt(rail_pad_pt)
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(max(8, int(round(style.name_font_size_pt * 0.92))))
        run.font.name = style.font_primary
        _apply_run_resume_color(run, style, "resume_text_color_emphasis")
        _set_run_character_spacing(run, style.name_letter_spacing_pt or 0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(style.name_space_after_pt)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add tagline (if visible).
    has_tagline = False
    vis = header.get("visibility", {})
    if vis.get("showTagline", True):
        tagline = (header.get("tagline") or "").strip()
        if tagline:
            has_tagline = True
            p = cell.add_paragraph()
            if rail_pad_pt:
                p.paragraph_format.left_indent = Pt(rail_pad_pt)
            for text, t_bold, t_italic, t_underline in parse_tagline_runs(tagline):
                run = p.add_run(text)
                run.font.size = Pt(style.tagline_font_size_pt)
                run.font.name = style.font_primary
                run.bold = t_bold
                run.italic = t_italic
                run.font.underline = (
                    WD_UNDERLINE.SINGLE if t_underline else WD_UNDERLINE.NONE
                )
                _apply_run_resume_color(run, style, "resume_text_color_primary")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(style.tagline_space_after_pt)
            _set_line_spacing_multiple(p, style.tagline_line_height)

    # Add contact rows based on visibility and ordering.
    contact_order = header.get("contactOrder", ["email", "phone", "location", "linkedin", "github", "portfolio"])

    field_map = {
        "email": header.get("email", ""),
        "phone": header.get("phone", ""),
        "location": header.get("location", ""),
        "linkedin": header.get("linkedin", ""),
        "github": header.get("github", ""),
        "portfolio": header.get("portfolio", ""),
    }
    visibility = header.get("visibility", {})
    visibility_map = {
        "email": "showEmail",
        "phone": "showPhone",
        "location": "showLocation",
        "linkedin": "showLinkedin",
        "github": "showGithub",
        "portfolio": "showPortfolio",
    }
    url_disp = resolve_contact_url_display(style_preferences)
    contact_idx = 0

    # Iterate over each contact field, adding to the cell with proper styling and added icon.
    for field_key in contact_order:
        if field_key not in field_map:
            continue
        val = field_map[field_key]
        vis_key = visibility_map.get(field_key)
        is_visible = visibility.get(vis_key, True) if vis_key else True
        if not val or not str(val).strip() or not is_visible:
            continue
        disp = format_contact_field_display(
            field_key,
            str(val).strip(),
            contact_url_display=url_disp,
        )
        p = cell.add_paragraph()
        if rail_pad_pt:
            p.paragraph_format.left_indent = Pt(rail_pad_pt)
        icon_pt = float(getattr(style, "sidebar_docx_contact_icon_size_pt", 0) or 0)
        icon_gap = float(getattr(style, "sidebar_docx_contact_icon_text_gap_pt", 4.5) or 0)
        png = (
            contact_rail_icon_png_bytes(template_slug, field_key)
            if icon_pt > 0
            else None
        )
        if png:
            try:
                r_img = p.add_run()
                r_img.add_picture(
                    BytesIO(png),
                    width=Pt(icon_pt),
                    height=Pt(icon_pt),
                )
                if icon_gap > 0:
                    n_nbsp = max(1, min(6, int(round(icon_gap / 2.5))))
                    r_sp = p.add_run("\u00a0" * n_nbsp)
                    r_sp.font.size = Pt(style.contact_font_size_pt)
                    r_sp.font.name = style.font_primary
            except (OSError, ValueError):
                pass
        run = p.add_run(disp)
        run.font.size = Pt(style.contact_font_size_pt)
        run.font.name = style.font_primary
        _apply_run_resume_color(run, style, "resume_text_color_primary")
        _set_run_character_spacing(run, style.contact_span_letter_spacing_pt)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        space_before = (
            style.contact_space_before_after_tagline_pt
            if has_tagline and contact_idx == 0
            else (style.contact_space_before_pt if contact_idx == 0 and not has_tagline else 0)
        )
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(1.5)
        _set_line_spacing_multiple(p, style.contact_line_height)
        contact_idx += 1

    if contact_idx and cell.paragraphs:
        base_after = float(style.contact_space_after_pt)
        trail = float(getattr(style, "sidebar_header_margin_bottom_pt", 0) or 0)
        cell.paragraphs[-1].paragraph_format.space_after = Pt(base_after + trail)
    elif cell.paragraphs:
        trail = float(getattr(style, "sidebar_header_margin_bottom_pt", 0) or 0)
        if trail > 0:
            lp = cell.paragraphs[-1]
            cur = lp.paragraph_format.space_after
            cur_pt = float(cur.pt) if cur else 0.0
            lp.paragraph_format.space_after = Pt(cur_pt + trail)


# --- Handle Single-column Centered Header ---
def _add_header(
    document: Document,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    style_preferences: dict | None = None,
) -> None:

    # Get header.
    header = resume_data.get("header", {})
    name = f"{header.get('first_name', '')} {header.get('last_name', '')}".strip()
    if name:
        p = document.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(style.name_font_size_pt)
        run.font.name = style.font_primary
        _apply_run_resume_color(run, style, "resume_text_color_emphasis")
        _set_run_character_spacing(run, style.name_letter_spacing_pt or 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(style.name_space_after_pt)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add tagline (if visible).
    has_tagline = False
    vis = header.get("visibility", {})
    if vis.get("showTagline", True):
        tagline = (header.get("tagline") or "").strip()
        if tagline:
            has_tagline = True
            p = document.add_paragraph()
            for text, t_bold, t_italic, t_underline in parse_tagline_runs(tagline):
                run = p.add_run(text)
                run.font.size = Pt(style.tagline_font_size_pt)
                run.font.name = style.font_primary
                run.bold = t_bold
                run.italic = t_italic
                run.font.underline = WD_UNDERLINE.SINGLE if t_underline else WD_UNDERLINE.NONE
                _apply_run_resume_color(run, style, "resume_text_color_primary")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(style.tagline_space_after_pt)
            _set_line_spacing_multiple(p, style.tagline_line_height)

    # Add contact line.
    contact_order = header.get("contactOrder", ["email", "phone", "location", "linkedin", "github", "portfolio"])
    field_map = {
        "email": header.get("email", ""), "phone": header.get("phone", ""), "location": header.get("location", ""),
        "linkedin": header.get("linkedin", ""), "github": header.get("github", ""), "portfolio": header.get("portfolio", ""),
    }

    # Get visibility.
    visibility = header.get("visibility", {})
    visibility_map = { "email": "showEmail", "phone": "showPhone", "location": "showLocation", "linkedin": "showLinkedin", "github": "showGithub", "portfolio": "showPortfolio", }
    
    # Get URL display preference.
    url_disp = resolve_contact_url_display(style_preferences)
    fields = []
    for field_key in contact_order:
        if field_key in field_map:
            val = field_map[field_key]
            vis_key = visibility_map.get(field_key)
            is_visible = visibility.get(vis_key, True) if vis_key else True
            if val and str(val).strip() and is_visible:
                fields.append(
                    format_contact_field_display(
                        field_key,
                        str(val).strip(),
                        contact_url_display=url_disp,
                    )
                )
    
    if fields:
        p = document.add_paragraph()
        run = p.add_run(" | ".join(fields))
        run.font.size = Pt(style.contact_font_size_pt)
        run.font.name = style.font_primary  # Georgia (contact line), not Times New Roman
        _apply_run_resume_color(run, style, "resume_text_color_primary")
        _set_run_character_spacing(run, style.contact_span_letter_spacing_pt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        space_before = (
            style.contact_space_before_after_tagline_pt
            if has_tagline
            else style.contact_space_before_pt
        )
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(style.contact_space_after_pt)
        _set_line_spacing_multiple(p, style.contact_line_height)
