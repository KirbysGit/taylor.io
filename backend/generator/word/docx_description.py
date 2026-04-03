# Experience / Project Description Blocks.

from __future__ import annotations

import re
from typing import Any, List, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.shared import Pt

from .docx_run_style import _apply_run_resume_color, _set_line_spacing_multiple
from .docx_styles import DocxStyleConfig

# Parse description into bullet/paragraph items.
def _parse_description_items(description: str) -> List[Tuple[str, str]]:
    if not description or not str(description).strip():
        return []

    # Split description into lines.
    lines = str(description).split("\n")
    items = []
    non_bullet_lines = []

    # Iterate over each line.
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # If the line starts with a bullet, add it to the bullet items.
        if stripped.startswith("•"):
            # Add any non-bullet lines to the items.
            if non_bullet_lines:
                items.append(("paragraph", "\n".join(non_bullet_lines)))
                non_bullet_lines = []
            # Remove the bullet from the line.
            bullet_text = re.sub(r"^•\s+", "", stripped)
            # Add the bullet text to the items.
            if bullet_text:
                items.append(("bullet", bullet_text))
        else:
            # Add the line to the non-bullet lines.
            non_bullet_lines.append(stripped)

    # Add any non-bullet lines to the items.
    if non_bullet_lines:
        items.append(("paragraph", "\n".join(non_bullet_lines)))
        
    # Return the items.
    return items


# Add description block.
def _add_description_block(doc: Any, style: DocxStyleConfig, description: str, indent_pt: float) -> None:

    # Parse description items.
    items = _parse_description_items(description)
    if not items:
        return

    # Initialize first flag.
    first = True
    for item_type, content in items:
        if not content or not str(content).strip():
            continue
        # Add paragraph.
        p = doc.add_paragraph()
        
        # Set alignment.
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Handle space before based on if first item or not.
        if first:
            if item_type == "bullet":
                p.paragraph_format.space_before = Pt(
                    style.description_block_space_before_pt + style.word_bullet_space_before_pt
                )
            else:
                p.paragraph_format.space_before = Pt(style.description_block_space_before_pt)
            first = False
        else:
            if item_type == "bullet":
                p.paragraph_format.space_before = Pt(style.word_bullet_space_before_pt)
            else:
                p.paragraph_format.space_before = Pt(0)

        # Handle bullet items.
        if item_type == "bullet":
            # Calculate bullet left padding.
            li_pad = float(getattr(style, "description_bullet_li_padding_left_pt", 2.0) or 0.0)
            hang = float(style.description_bullet_hang_pt)
            bullet_left_pt = indent_pt + style.description_bullet_indent_pt

            # Calculate text left padding.
            text_left_pt = bullet_left_pt + hang + li_pad

            # Add tab stop.
            p.paragraph_format.left_indent = Pt(text_left_pt)
            p.paragraph_format.first_line_indent = Pt(-(hang + li_pad))
            p.paragraph_format.tab_stops.add_tab_stop(
                Pt(text_left_pt),
                WD_TAB_ALIGNMENT.LEFT,
            )

            # Set space after and line spacing rule.
            p.paragraph_format.space_after = Pt(style.word_bullet_space_after_pt)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # Set alignment.
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_b = p.add_run("•")
            r_b.font.size = Pt(style.description_font_size_pt)
            r_b.font.name = style.font_primary
            _apply_run_resume_color(r_b, style, "resume_text_color_primary")
            r_tab = p.add_run("\t")
            r_tab.font.size = Pt(style.description_font_size_pt)
            r_tab.font.name = style.font_primary
            run = p.add_run(str(content).strip())
        else:
            # Set left indent and space after.
            p.paragraph_format.left_indent = Pt(indent_pt)
            p.paragraph_format.space_after = Pt(style.description_paragraph_space_pt)
            run = p.add_run(str(content).strip())
            _set_line_spacing_multiple(p, style.prose_line_height)
        
        # Set font size, name and color.
        run.font.size = Pt(style.description_font_size_pt)
        run.font.name = style.font_primary
        _apply_run_resume_color(run, style, "resume_text_color_primary")
