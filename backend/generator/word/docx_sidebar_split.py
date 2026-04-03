# Sidebar-split layout: table row, rail column, main column, pagination heuristics.

# In this file,


from __future__ import annotations

from io import BytesIO
from dataclasses import replace
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from ..layouts.sidebar_split import sidebar_main_column_order, sidebar_rail_section_order
from ..shared.dates import format_date_range
from ..shared.skills import skills_group_ordered
from .docx_header import _add_sidebar_rail_header
from .docx_layout import _apply_section_title_bottom_border
from .docx_run_style import _apply_run_resume_color, _hex_rgb_6_for_word, _set_line_spacing_multiple
from .docx_sections import (
    _render_docx_experience_section,
    _render_docx_projects_section,
    _render_docx_summary_section,
)
from .docx_styles import DocxStyleConfig


# --- Handle Set Cell Margins Inches ---
def _set_cell_margins_inches(cell: Any, *, top: float = 0, left: float = 0, bottom: float = 0, right: float = 0) -> None:
    """Word tcMar in twips (dxa): inset cell text from cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in list(tcPr.findall(qn("w:tcMar"))):
        tcPr.remove(old)
    if top <= 0 and left <= 0 and bottom <= 0 and right <= 0:
        return
    tcMar = OxmlElement("w:tcMar")
    for tag, inches in (
        ("w:top", top),
        ("w:left", left),
        ("w:bottom", bottom),
        ("w:right", right),
    ):
        if inches <= 0:
            continue
        el = OxmlElement(tag)
        el.set(qn("w:w"), str(int(round(float(inches) * 1440))))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


# --- Handle Sidebar DOCX: stretch table row to page height only when content likely fits on one page. ---
# Word may still paginate oddly if the row splits across pages (tall fixed row + overflow).
_SIDEBAR_DOCX_FILL_ROW_SLACK_TWIPS = 720  # heuristic one-pager: room for pagination / sectPr
_SIDEBAR_DOCX_FILL_ROW_SLACK_TWIPS_DOCX_MAX_1 = (
    720  # ~0.5in under full body + stub para before sectPr - fewer phantom pages than 240
)
_SIDEBAR_DOCX_AVG_LINE_PT = 12.75
_SIDEBAR_DOCX_ONE_PAGE_LINES_FRAC = 0.78  # conservative; underestimate risk = blank page again

# --- Handle Sidebar DOCX Wrapped Line Count ---
def _sidebar_docx_wrapped_line_count(text: str, chars_per_line: float = 58.0) -> float:
    # Get text.
    t = (text or "").strip()

    # Check if text is empty.
    if not t:
        return 0.0

    # Initialize total.
    total = 0.0

    # Iterate over each line.
    for seg in t.split("\n"):
        seg = seg.strip()
        if not seg:
            # Add 0.35 to total for empty line.
            total += 0.35
            continue

        # Add 1.0 to total for each segment.
        total += max(1.0, len(seg) / chars_per_line)
    return total

# --- Handle Estimate Sidebar DOCX Main Column Lines ---
def _estimate_sidebar_docx_main_column_lines(resume_data: Dict[str, Any]) -> float:
    # Initialize lines.
    lines = 0.0

    # Iterate over each key in sidebar main column order.
    for key in sidebar_main_column_order(resume_data):
        if key == "summary":
            # Get summary.
            summ = (resume_data.get("summary") or {}).get("summary") or ""
            if not str(summ).strip():
                continue
            lines += 3.4
            lines += _sidebar_docx_wrapped_line_count(str(summ), 56.0)
        elif key == "experience":
            # Get experience.
            exps = resume_data.get("experience") or []
            if not exps:
                continue
            lines += 3.6

            for exp in exps:
                if not isinstance(exp, dict):
                    exp = {}
                lines += 2.8
                desc = str(exp.get("description") or "")
                lines += _sidebar_docx_wrapped_line_count(desc, 56.0) * 1.08
                lines += 1.35
        elif key == "projects":
            # Get projects.
            projs = resume_data.get("projects") or []
            if not projs:
                continue
            lines += 3.6

            # Iterate over each project.
            for proj in projs:
                if not isinstance(proj, dict):
                    proj = {}
                lines += 3.2
                desc = str(proj.get("description") or "")
                lines += _sidebar_docx_wrapped_line_count(desc, 56.0)
                lines += 0.9

    return lines

# --- Handle Estimate Sidebar DOCX Rail Lines ---
def _estimate_sidebar_docx_rail_lines(resume_data: Dict[str, Any]) -> float:
    # Initialize lines.
    lines = 6.5
    header = resume_data.get("header") or {}
    vis = header.get("visibility") or {}
    if vis.get("showTagline", True) and (header.get("tagline") or "").strip():
        lines += 1.6

    # Get contact order.
    order = header.get("contactOrder", ["email", "phone", "location", "linkedin", "github", "portfolio"])
    field_map = {
        "email": header.get("email", ""),
        "phone": header.get("phone", ""),
        "location": header.get("location", ""),
        "linkedin": header.get("linkedin", ""),
        "github": header.get("github", ""),
        "portfolio": header.get("portfolio", ""),
    }
    vis_map = {
        "email": "showEmail",
        "phone": "showPhone",
        "location": "showLocation",
        "linkedin": "showLinkedin",
        "github": "showGithub",
        "portfolio": "showPortfolio",
    }
    n_contact = 0

    # Iterate over each contact field.
    for fk in order:
        if fk not in field_map:
            continue
        val = str(field_map.get(fk) or "").strip()
        if not val:
            continue
        vk = vis_map.get(fk)
        if vk and not vis.get(vk, True):
            continue
        n_contact += 1
    lines += max(0.0, n_contact - 1) * 1.05 + (1.0 if n_contact else 0.0)

    # Iterate over each key in sidebar rail section order.
    for key in sidebar_rail_section_order(resume_data):
        if key == "skills":
            skills = resume_data.get("skills") or []
            if not skills:
                continue
            lines += 3.0
            cat_order = resume_data.get("skillsCategoryOrder") or []
            groups = skills_group_ordered(skills, cat_order if cat_order else None)
            for category, names in groups:
                if not names:
                    continue
                if category:
                    lines += 1.15
                # Narrow column: names wrap more often
                joined_len = sum(len(str(n)) for n in names) + max(0, len(names) - 1) * 2
                lines += max(1.0, joined_len / 34.0)
                lines += 0.35
        elif key == "education":
            edu_list = resume_data.get("education") or []
            if not edu_list:
                continue
            lines += 3.0
            for edu in edu_list:
                if not isinstance(edu, dict):
                    edu = {}
                subs = edu.get("subsections") or {}
                if not isinstance(subs, dict):
                    subs = {}
                lines += 5.8 + min(8.0, len(subs)) * 1.8
    return lines

# --- Handle Sidebar DOCX Row Likely Fits One Page ---
def _sidebar_docx_row_likely_fits_one_page(sect: Any, resume_data: Dict[str, Any]) -> bool:

    # Get body height in inches.
    body_h_in = float(sect.page_height.inches) - float(sect.top_margin.inches) - float(
        sect.bottom_margin.inches
    )

    # Get usable lines.
    usable_lines = (max(1.0, body_h_in) * 72.0) / _SIDEBAR_DOCX_AVG_LINE_PT
    cap = usable_lines * _SIDEBAR_DOCX_ONE_PAGE_LINES_FRAC

    # Get main and rail lines.
    main_l = _estimate_sidebar_docx_main_column_lines(resume_data)
    rail_l = _estimate_sidebar_docx_rail_lines(resume_data)

    # Return if the main or rail lines are less than or equal to the capacity.
    return max(main_l, rail_l) <= cap


# --- Handle Finalize Sidebar Split Table ---
def _finalize_sidebar_split_table(
    tbl: Any,
    sect: Any,
    rail_in: float,
    main_in: float,
    *,
    fill_row_to_page: bool = False,
    fill_slack_twips: int = _SIDEBAR_DOCX_FILL_ROW_SLACK_TWIPS,
    fill_row_height_in: float | None = None,
) -> None:
    # Get usable inches.
    usable_in = (
        float(sect.page_width.inches)
        - float(sect.left_margin.inches)
        - float(sect.right_margin.inches)
    )

    # Get usable twips.
    usable_twips = int(round(max(0.5, usable_in) * 1440))

    # Get rail and main twips.
    rail_twips = int(round(max(0.25, float(rail_in)) * 1440))
    main_twips = max(int(round(float(main_in) * 1440)), 1)
    # Keep grid exact to printable width so Word does not rescale columns.
    drift = usable_twips - (rail_twips + main_twips)
    if drift != 0:
        main_twips += drift

    # Get table element.
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)

    # Get table width.
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(usable_twips))

    # Get table layout.
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    # Get table grid.
    tbl_grid = tbl_el.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        ins = len(tbl_el)
        for i, child in enumerate(tbl_el):
            if child.tag == qn("w:tr"):
                ins = i
                break
        tbl_el.insert(ins, tbl_grid)
    for old_gc in list(tbl_grid.findall(qn("w:gridCol"))):
        tbl_grid.remove(old_gc)
    for w_tw in (rail_twips, main_twips):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w_tw))
        tbl_grid.append(gc)

    # Iterate over each column.
    for ci, w_tw in enumerate((rail_twips, main_twips)):
        tc = tbl.rows[0].cells[ci]._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.insert(0, tc_w)
        tc_w.set(qn("w:type"), "dxa")
        tc_w.set(qn("w:w"), str(w_tw))

    # Get table row.
    tr = tbl.rows[0]._tr
    tr_pr = tr.find(qn("w:trPr"))
    if tr_pr is not None:
        for old_th in list(tr_pr.findall(qn("w:trHeight"))):
            tr_pr.remove(old_th)
        if not fill_row_to_page and len(list(tr_pr)) == 0:
            tr.remove(tr_pr)
            tr_pr = None

    # If filling row to page, add table row properties.
    if fill_row_to_page:
        if tr_pr is None:
            tr_pr = OxmlElement("w:trPr")
            tr.insert(0, tr_pr)

        # Get body height in inches.
        body_h_in = (
            float(sect.page_height.inches)
            - float(sect.top_margin.inches)
            - float(sect.bottom_margin.inches)
        )

        # Get body twips.
        body_twips = int(round(max(1.0, body_h_in) * 1440))

        # If fill row height in is not None and greater than 0, set target to the fill row height in.
        if fill_row_height_in is not None and float(fill_row_height_in) > 0:
            target = max(1000, int(round(float(fill_row_height_in) * 1440)))
            target = min(target, body_twips)
        else:
            slack = max(0, int(fill_slack_twips))
            target = max(1000, body_twips - slack)
        tr_h = OxmlElement("w:trHeight")
        tr_h.set(qn("w:val"), str(target))
        tr_h.set(qn("w:hRule"), "exact")
        tr_pr.append(tr_h)

# --- Handle Apply Sidebar Rail Cell Appearance ---
def _apply_sidebar_rail_cell_appearance(cell: Any, style: DocxStyleConfig) -> None:
    # Get table cell element.
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for tag in ("w:shd", "w:tcBorders"):
        for old in list(tcPr.findall(qn(tag))):
            tcPr.remove(old)

    # Get fill color.
    fill = _hex_rgb_6_for_word(getattr(style, "sidebar_bg", "") or "")
    if fill:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    # Get border color.
    border_c = _hex_rgb_6_for_word(
        getattr(style, "sidebar_border_color", None) or "#b8c2ce"
    )
    # Get border width.
    border_w = float(getattr(style, "sidebar_border_width_pt", 0.75) or 0.75)
    sz = max(2, min(96, int(round(border_w * 8))))

    # If border color is not empty, add table cell borders.
    if border_c:
        tb = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "nil")
            tb.append(b)
        br = OxmlElement("w:right")
        br.set(qn("w:val"), "single")
        br.set(qn("w:sz"), str(sz))
        br.set(qn("w:space"), "0")
        br.set(qn("w:color"), border_c)
        tb.append(br)
        tcPr.append(tb)

# --- Handle Add Sidebar DOCX Section Title ---
def _add_section_title_sidebar_rail(
    doc: Any,
    title: str,
    style: DocxStyleConfig,
    *,
    left_indent_pt: float = 0,
    section_stack_margin_top_pt: float = 0,
) -> None:
    # Add section title paragraph.
    p = doc.add_paragraph()
    if left_indent_pt:
        p.paragraph_format.left_indent = Pt(left_indent_pt)

    # Add section title run.
    run = p.add_run(title)
    sz = max(8, int(round(style.section_title_font_size_pt * 0.9)))
    run.font.size = Pt(sz)
    run.font.name = style.font_primary
    _apply_run_resume_color(run, style, "resume_text_color_emphasis")
    p.paragraph_format.space_before = Pt(
        style.section_title_space_before_pt * 0.85 + section_stack_margin_top_pt
    )
    p.paragraph_format.space_after = Pt(style.section_title_space_after_pt * 0.85)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _apply_section_title_bottom_border(p)

# --- Handle Render Sidebar DOCX Education Rail Section ---
def _render_docx_education_rail(
    doc: Any,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
    *,
    section_stack_margin_top_pt: float = 0,
) -> bool:
    # Get education.
    education = resume_data.get("education") or []
    rendered_any = False
    school_sz = max(8, int(round(style.school_name_font_size_pt * 0.92)))
    meta_sz = max(7, int(round(style.school_meta_font_size_pt * 0.88)))
    deg_sz = max(8, int(round(style.education_degree_line_font_size_pt * 1.05)))
    minor_sz = max(7, float(getattr(style, "education_rail_minor_font_size_pt", 9.0) or 9.0))
    hl_sz = max(7, int(round(style.highlight_font_size_pt)))
    gpa_after = float(getattr(style, "education_rail_gpa_space_after_pt", 6.0) or 6.0)
    hl_before = float(getattr(style, "education_rail_highlights_space_before_pt", 8.0) or 8.0)
    entry_tail_gap = float(
        getattr(style, "education_entry_margin_bottom_pt", 8.0) or 8.0
    )

    # Iterate over each education entry.
    for edu in education:
        if not isinstance(edu, dict):
            edu = {}
        school_raw = (edu.get("school") or "").strip()
        degree = (edu.get("degree") or "").strip()
        discipline = (edu.get("discipline") or "").strip()
        minor = (edu.get("minor") or "").strip()
        location = (edu.get("location") or "").strip()
        gpa = (edu.get("gpa") or "").strip()
        start_raw = edu.get("start_date") or edu.get("startDate") or ""
        end_raw = edu.get("end_date") or edu.get("endDate") or ""
        date_range = format_date_range(start_raw, end_raw, edu.get("current", False))
        minor_line = f"Minor in {minor}" if minor else ""
        has_major = bool(degree or discipline)
        subs = edu.get("subsections") or {}
        if not isinstance(subs, dict):
            subs = {}
        has_body = bool(
            school_raw or date_range or location or has_major or minor_line or gpa
        ) or bool(subs)

        # If there is no body, continue.
        if not has_body:
            continue

        # If no education has been rendered yet, add the section title.
        if not rendered_any:
            _add_section_title_sidebar_rail(
                doc,
                section_labels.get("education", defaults["education"]),
                style,
                left_indent_pt=indent,
                section_stack_margin_top_pt=section_stack_margin_top_pt,
            )
            rendered_any = True

        # If there is a school raw, add the school paragraph.
        if school_raw:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(indent)
            p.paragraph_format.space_after = Pt(style.school_line_space_after_pt)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(school_raw)
            r.bold = True
            r.font.size = Pt(school_sz)
            r.font.name = style.font_primary
            _apply_run_resume_color(r, style, "resume_text_color_emphasis")
        if date_range:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(indent)
            p.paragraph_format.space_after = Pt(1.25)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(date_range)
            r.font.size = Pt(meta_sz)
            r.font.name = style.font_secondary
            r.italic = True
            _apply_run_resume_color(r, style, "resume_text_color_secondary")
        if location:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(indent)
            p.paragraph_format.space_after = Pt(1.25)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(location)
            r.font.size = Pt(meta_sz)
            r.font.name = style.font_secondary
            r.italic = True
            _apply_run_resume_color(r, style, "resume_text_color_secondary")
        if degree and discipline:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(indent)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r1 = p.add_run(f"{degree} in ")
            r1.font.size = Pt(deg_sz)
            r1.font.name = style.font_secondary
            _apply_run_resume_color(r1, style, "resume_text_color_primary")
            r2 = p.add_run(discipline)
            r2.font.size = Pt(deg_sz)
            r2.font.name = style.font_secondary
            _apply_run_resume_color(r2, style, "resume_text_color_primary")
        elif degree:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(indent)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(degree)
            r.font.size = Pt(deg_sz)
            r.font.name = style.font_secondary
            _apply_run_resume_color(r, style, "resume_text_color_primary")
        elif discipline:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(indent)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(discipline)
            r.font.size = Pt(deg_sz)
            r.font.name = style.font_secondary
            _apply_run_resume_color(r, style, "resume_text_color_primary")
        if minor_line:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(indent)
            p.paragraph_format.space_after = Pt(1.25)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(minor_line)
            r.font.size = Pt(minor_sz)
            r.font.name = style.font_secondary
            _apply_run_resume_color(r, style, "resume_text_color_primary")
        if gpa:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(indent)
            p.paragraph_format.space_after = Pt(gpa_after)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(f"GPA: {gpa}")
            r.font.size = Pt(style.school_gpa_font_size_pt)
            r.font.name = style.font_secondary
            r.italic = True
            _apply_run_resume_color(r, style, "resume_text_color_secondary")

        hl_gap_pt = float(getattr(style, "highlights_gap_pt", 4.5) or 4.5)
        hl_block_i = 0
        for sub_title, sub_content in subs.items():
            st = str(sub_title).strip() if sub_title is not None else ""
            sc = str(sub_content).strip() if sub_content is not None else ""
            if not st and not sc:
                continue
            block_top = hl_before if hl_block_i == 0 else hl_gap_pt
            if st:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(indent)
                p.paragraph_format.space_before = Pt(block_top)
                p.paragraph_format.space_after = Pt(1.0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                r = p.add_run(st)
                r.bold = True
                r.font.size = Pt(hl_sz)
                r.font.name = style.font_primary
                _apply_run_resume_color(r, style, "resume_text_color_highlight_title")
            if sc:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(indent)
                p.paragraph_format.space_before = Pt(0 if st else block_top)
                # No space_after: PDF rail uses flex gap between .hl blocks, not margin on body.
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                r = p.add_run(sc)
                r.font.size = Pt(hl_sz)
                r.font.name = style.font_primary
                _apply_run_resume_color(r, style, "resume_text_color_primary")

            hl_block_i += 1

        if doc.paragraphs:
            doc.paragraphs[-1].paragraph_format.space_after = Pt(entry_tail_gap)

    return rendered_any

# --- Handle Render Sidebar DOCX Skills Rail Section ---
def _render_docx_skills_rail_section(
    doc: Any,
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    indent: float,
    section_labels: Dict[str, Any],
    defaults: Dict[str, str],
    *,
    section_stack_margin_top_pt: float = 0,
) -> bool:
    # Get skills.
    skills = resume_data.get("skills") or []
    category_order = resume_data.get("skillsCategoryOrder") or []
    groups = skills_group_ordered(skills, category_order if category_order else None)
    if not groups:
        return False

    # Add section title.
    _add_section_title_sidebar_rail(
        doc,
        section_labels.get("skills", defaults["skills"]),
        style,
        left_indent_pt=indent,
        section_stack_margin_top_pt=section_stack_margin_top_pt,
    )

    # Get skill category and name font sizes.
    cat_sz = max(8, int(round(style.skill_category_font_size_pt * 0.95)))
    name_sz = max(8, int(round(style.skill_names_font_size_pt * 0.95)))

    # Iterate over each group.
    for category, names in groups:
        if not names:
            continue
        if category:
            # Add skill category paragraph.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(indent)
            p.paragraph_format.space_after = Pt(2.0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(str(category))
            r.bold = True
            r.font.size = Pt(cat_sz)
            r.font.name = style.font_primary
            _apply_run_resume_color(r, style, "resume_text_color_emphasis")

        # Add skill lines paragraph.
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Pt(indent)
        p2.paragraph_format.space_after = Pt(style.skill_line_space_pt)
        _set_line_spacing_multiple(p2, style.skill_line_height)
        r2 = p2.add_run(", ".join(names))
        r2.font.size = Pt(name_sz)
        r2.font.name = style.font_primary
        _apply_run_resume_color(r2, style, "resume_text_color_primary")

    return True

# --- Handle Remove Singleton Empty Paragraph ---
def _remove_singleton_empty_paragraph(cell: Any) -> None:
    # Get table cell paragraphs.
    paras = cell.paragraphs
    # If there are not exactly one paragraph, return.
    if len(paras) != 1:
        return
    if paras[0].text.strip():
        return
    el = paras[0]._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

# --- Handle Strip Document Leading Empty Paragraph ---
def _strip_document_leading_empty_paragraph(doc: Document) -> None:
    # If there are no paragraphs, return.
    if not doc.paragraphs:
        return
    p0 = doc.paragraphs[0]
    # If the first paragraph is not empty, return.
    if p0.text.strip():
        return
    el = p0._element
    # If the parent is not None, remove the element.
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

# --- Handle Sidebar DOCX Tail Paragraph Before SectPr ---
def _sidebar_docx_tail_paragraph_before_sectpr(doc: Document) -> None:
    # Get body.
    body = doc.element.body
    kids = list(body)
    if len(kids) < 2 or kids[-1].tag != qn("w:sectPr"):
        return
    if kids[-2].tag == qn("w:p"):
        return
    idx = len(kids) - 1
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "20")
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)
    p.append(pPr)

    # Add run.
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "2")
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), "2")
    rPr.append(sz)
    rPr.append(sz_cs)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = "\u200b"
    r.append(t)
    p.append(r)
    body.insert(idx, p)


# --- Handle Build Sidebar DOCX Split Document ---
def _build_docx_sidebar_split_document(
    resume_data: Dict[str, Any],
    style: DocxStyleConfig,
    style_preferences: dict | None = None,
    *,
    template_slug: str = "sidebar",
    docx_max_pages: int | None = None,
) -> bytes:

    # Initialize document.
    doc = Document()
    # Get section.
    sect = doc.sections[0]

    # Match PDF: Playwright margin is 0 for sidebar; resume is a full Letter box and
    # preview.css applies margin_top / pad_x / margin_bottom inside each column.
    sect.top_margin = Inches(0)
    sect.bottom_margin = Inches(0)
    sect.left_margin = Inches(0)
    sect.right_margin = Inches(0)

    # Word reserves ~0.5" for header/footer in w:pgMar even when empty; clear it so the
    # table can align with the physical page like the PDF (Playwright margin 0).
    sect.header_distance = Inches(0)
    sect.footer_distance = Inches(0)

    # Strip document leading empty paragraph.
    _strip_document_leading_empty_paragraph(doc)

    # Get usable points.
    usable_pt = (
        float(sect.page_width.pt)
        - float(sect.left_margin.pt)
        - float(sect.right_margin.pt)
    )
    usable_in = usable_pt / 72.0
    rail_in = float(style.sidebar_width_in)
    main_in = max(1.25, usable_in - rail_in)

    # Add table and set columns.
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    try:
        tbl.allow_autofit = False
    except AttributeError:
        pass
    tbl.columns[0].width = Inches(rail_in)
    tbl.columns[1].width = Inches(main_in)

    # Get left and right cells.
    left_cell = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]
    _remove_singleton_empty_paragraph(left_cell)
    _remove_singleton_empty_paragraph(right_cell)

    # PDF .resume-aside: padding = margin_top, pad_x, margin_bottom, pad_x (shading still
    # fills the rail). No page margin + no left tcMar so the rail is flush to the paper edge;
    # text uses paragraph left_indent from sidebar_pad_x_in (header + section renders).
    pad_x = float(style.sidebar_pad_x_in)
    mt = float(style.margin_top_in)
    mb = float(style.margin_bottom_in)
    mr = float(style.margin_right_in)
    _set_cell_margins_inches(
        left_cell, top=mt, bottom=mb, left=0, right=pad_x
    )
    _apply_sidebar_rail_cell_appearance(left_cell, style)
    main_left_pad = float(style.sidebar_body_pad_left_in)
    _set_cell_margins_inches(
        right_cell, top=mt, bottom=mb, left=main_left_pad, right=mr
    )

    # Add sidebar rail header.
    _add_sidebar_rail_header(
        left_cell,
        resume_data,
        style,
        style_preferences,
        template_slug=template_slug,
    )

    # Get section labels and defaults.
    section_labels = resume_data.get("sectionLabels", {})
    defaults = {
        "summary": "Professional Summary",
        "education": "Education",
        "experience": "Experience",
        "projects": "Projects",
        "skills": "Skills",
    }

    # Get rail indent.
    rail_indent = float(style.sidebar_pad_x_in) * 72.0

    # Iterate over each rail block.
    rail_blocks = 0
    for key in sidebar_rail_section_order(resume_data):
        stack_m = (
            float(getattr(style, "sidebar_rail_section_stack_margin_top_pt", 0) or 0)
            if rail_blocks
            else 0.0
        )

        # If the key is skills, render the skills rail section.
        if key == "skills":
            if _render_docx_skills_rail_section(
                left_cell,
                resume_data,
                style,
                rail_indent,
                section_labels,
                defaults,
                section_stack_margin_top_pt=stack_m,
            ):
                rail_blocks += 1

        # If the key is education, render the education rail section.
        elif key == "education":
            if _render_docx_education_rail(
                left_cell,
                resume_data,
                style,
                rail_indent,
                section_labels,
                defaults,
                section_stack_margin_top_pt=stack_m,
            ):
                rail_blocks += 1
    
    # Get padding and tab inches.
    inner_main_in = max(1.0, main_in - main_left_pad - mr)
    tab_main = max(1.25, inner_main_in - 0.01)
    
    # Get main style.
    style_main = replace(
        style,
        two_column_tab_in=tab_main,
        summary_text_padding_left_pt=float(style.sidebar_summary_text_padding_left_pt),
        summary_first_line_indent_pt=float(style.sidebar_summary_first_line_indent_pt),
    )
    main_indent = float(style.sidebar_section_content_indent_pt)

    # Iterate over each main column block.
    for key in sidebar_main_column_order(resume_data):
        # If the key is summary, render the summary section.
        if key == "summary":
            _render_docx_summary_section(
                right_cell, resume_data, style_main, main_indent, section_labels, defaults
            )
        # If the key is experience, render the experience section.
        elif key == "experience":
            _render_docx_experience_section(
                right_cell, resume_data, style_main, main_indent, section_labels, defaults
            )
        # If the key is projects, render the projects section.
        elif key == "projects":
            _render_docx_projects_section(
                right_cell,
                resume_data,
                style_main,
                main_indent,
                section_labels,
                defaults,
                sidebar_main_column=True,
            )

    # If the document max pages is 1, set fill row to True and slack to the maximum slack twips.
    if docx_max_pages == 1:
        fill_row = True
        slack = _SIDEBAR_DOCX_FILL_ROW_SLACK_TWIPS_DOCX_MAX_1
    # Otherwise, set fill row to the likely fits one page and slack to the default slack twips.
    else:
        fill_row = _sidebar_docx_row_likely_fits_one_page(sect, resume_data)
        slack = _SIDEBAR_DOCX_FILL_ROW_SLACK_TWIPS
    row_h_in = float(getattr(style, "sidebar_docx_fill_row_height_in", 0.0) or 0.0)
    _finalize_sidebar_split_table(
        tbl,
        sect,
        rail_in,
        main_in,
        fill_row_to_page=fill_row,
        fill_slack_twips=slack,
        fill_row_height_in=row_h_in if row_h_in > 0 else None,
    )
    if fill_row:
        _sidebar_docx_tail_paragraph_before_sectpr(doc)

    # Save document to buffer and return the buffer.
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
